"""
cctv_app.py  (v7.6 — index page: real Wasabi check per bookmark on load; card status driven by actual cloud presence)
================================================
Changes from v6.4:
- Wasabi cloud upload integration across all 3 pipelines (SYP, RMBC, FOI)
- On form load: checks Wasabi for matching footage (video + JSON) for bookmark
- If footage found: officer chooses Download / Cloud / Both
- If not found: download only with deferred [D] upload warning
- Cloud upload timestamps recorded; upload failures shown but never block download
- WASABI_ACCESS_KEY, WASABI_SECRET_KEY, WASABI_BUCKET, WASABI_REGION in .env
- # TODO: cross-pipeline handover logic review (deferred — flagged in code)

Changes from v6.3:
- Anthropic API dependency removed entirely — Ollama only
- ANTHROPIC_API_KEY, CLAUDE_API_URL, CLAUDE_MODEL, AI_BACKEND constants removed
- No data leaves the building — all AI inference runs locally

Changes from v6.2:
- Prompt rewritten as strict transcription tool instructions — no creativity, no rephrasing
- temperature:0 added to Ollama payload — deterministic output, stops model improvising
- System prompt tightened to match

Changes from v6.1:
- AI_BACKEND config: set to "anthropic" (default) or "ollama" in .env
- Ollama support: calls local Ollama API (http://localhost:11434) when AI_BACKEND=ollama
- No data leaves the building when using Ollama — removes third-party processor concern
- OLLAMA_MODEL defaults to llama3 but is configurable via .env
- ANTHROPIC_API_KEY no longer required when AI_BACKEND=ollama

Changes from v6:
- Gmail SMTP email feature: send statement as .docx attachment
- Recipient built from name input + domain dropdown
- Domains: @rotherham.gov.uk / @southyorkshire.police.uk
- Email button appears in success box after statement is generated
- Temp file used to hold docx between generation and send — no re-generation
- Gmail credentials stored in .env file (GMAIL_USER, GMAIL_APP_PASSWORD)

Changes from v5:
- Site reference auto-grabbed from server hostname (socket.gethostname())
- Removed Wasabi Cloud Storage; added Raw Data (USB/Hard Drive) option
- Added Flare Reference field to incident reference section
- Camera mounting position pre-filled as Lamp Post with number prompt
- In-person handover: custody disclaimer added to form and statement
- Complete Later: wet ink workflow paragraph added to statement
- FOI pipeline: separate audit/registry Word document generated
- Visual improvements to Word document output
- DEMS only for electronic transfers (Wasabi removed throughout)

Install:
    python3 -m venv /opt/CCTV_Statement/venv
    /opt/CCTV_Statement/venv/bin/pip install flask python-docx requests

.env file must contain:
    OLLAMA_MODEL=llama3.1
    OLLAMA_HOST=http://localhost:11434  # optional, defaults to localhost
    GMAIL_USER=rmbcvms@gmail.com
    GMAIL_APP_PASSWORD=your-16-char-app-password
    WASABI_ACCESS_KEY=your-access-key
    WASABI_SECRET_KEY=your-secret-key
    WASABI_BUCKET=cctvserver
    WASABI_REGION=eu-west-1

Run:
    /opt/CCTV_Statement/venv/bin/python3 cctv_app.py

Access:
    http://0.0.0.0:5000
"""

import os, io, uuid, socket, sqlite3, hashlib, secrets, requests, smtplib, tempfile
from datetime import datetime, timezone
from functools import wraps
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from flask import Flask, render_template_string, request, redirect, url_for, session, send_file, jsonify
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)
app.secret_key = secrets.token_hex(32)

NX_SQLITE          = "/opt/networkoptix/mediaserver/var/mserver.sqlite"
GMAIL_USER         = os.environ.get("GMAIL_USER", "")
GMAIL_APP_PASSWORD = os.environ.get("GMAIL_APP_PASSWORD", "")
OLLAMA_HOST        = os.environ.get("OLLAMA_HOST", "http://localhost:11434")
OLLAMA_MODEL       = os.environ.get("OLLAMA_MODEL", "llama3.1")
WASABI_ACCESS_KEY  = os.environ.get("WASABI_ACCESS_KEY", "")
WASABI_SECRET_KEY  = os.environ.get("WASABI_SECRET_KEY", "")
WASABI_BUCKET      = os.environ.get("WASABI_BUCKET", "cctvserver")
WASABI_REGION      = os.environ.get("WASABI_REGION", "eu-west-1")

# Temp file store: token → (filepath, filename)
_TEMP_DOCS: dict = {}

# ── Auto site reference from hostname ─────────────────────────────────────────
def get_site_ref():
    try:
        hn = socket.gethostname().upper().split(".")[0]
        return hn
    except Exception:
        return "RMBC-UNKNOWN"

SITE_REF = get_site_ref()

USERS = {
    "dane.plant": {"hash": hashlib.sha256(b"Cctv2026!").hexdigest(),  "name": "Dane Plant",    "role": "CCTV Engineer",  "initials": "DP"},
    "admin":      {"hash": hashlib.sha256(b"Admin2026!").hexdigest(), "name": "Administrator", "role": "CCTV Manager",   "initials": "AD"},
}

LOCATIONS = [
    "Rawmarsh Police Station, Green Lane, Rawmarsh, Rotherham, S62 6JU",
    "Riverside House, Main Street, Rotherham, S60 1AE",
]

# ── Auth ──────────────────────────────────────────────────────────────────────

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if "username" not in session:
            return redirect(url_for("login"))
        return f(*args, **kwargs)
    return decorated

def check_password(username, password):
    user = USERS.get(username)
    if not user: return False
    return user["hash"] == hashlib.sha256(password.encode()).hexdigest()

# ── SQLite ────────────────────────────────────────────────────────────────────

def get_bookmarks():
    try:
        con = sqlite3.connect(NX_SQLITE)
        con.row_factory = sqlite3.Row
        cur = con.cursor()
        cur.execute("SELECT record_id, start_time, duration, name, description, created FROM bookmarks WHERE duration > 0 ORDER BY start_time DESC LIMIT 50")
        rows = [dict(r) for r in cur.fetchall()]
        con.close()
        for r in rows:
            s = datetime.fromtimestamp(r["start_time"] / 1000, tz=timezone.utc)
            e = datetime.fromtimestamp((r["start_time"] + r["duration"]) / 1000, tz=timezone.utc)
            d = r["duration"] / 1000
            # Seconds omitted intentionally — keyframe misalignment means exact
            # seconds from bookmark metadata are unreliable. HH:MM is sufficient
            # for all legal purposes and avoids misleading precision in court.
            r["start_fmt"]    = s.strftime("%d/%m/%Y %H:%M")
            r["end_fmt"]      = e.strftime("%d/%m/%Y %H:%M")
            r["duration_fmt"] = f"{int(d//60)}m {int(d%60)}s"
        return rows
    except Exception as e:
        print(f"SQLite error: {e}"); return []

def get_bookmark(record_id):
    try:
        con = sqlite3.connect(NX_SQLITE)
        con.row_factory = sqlite3.Row
        cur = con.cursor()
        cur.execute("SELECT record_id, hex(guid) as guid, start_time, duration, name, description, created FROM bookmarks WHERE record_id=?", (record_id,))
        r = dict(cur.fetchone()); con.close()
        s = datetime.fromtimestamp(r["start_time"] / 1000, tz=timezone.utc)
        e = datetime.fromtimestamp((r["start_time"] + r["duration"]) / 1000, tz=timezone.utc)
        d = r["duration"] / 1000
        r["start_dt"] = s; r["end_dt"] = e
        # Seconds omitted intentionally — keyframe misalignment means exact
        # seconds from bookmark metadata are unreliable. HH:MM is sufficient
        # for all legal purposes and avoids misleading precision in court.
        r["start_fmt"]    = s.strftime("%d/%m/%Y %H:%M")
        r["end_fmt"]      = e.strftime("%d/%m/%Y %H:%M")
        hours = int(d // 3600)
        mins  = int((d % 3600) // 60)
        secs  = int(d % 60)
        if hours > 0:
            r["duration_fmt"] = f"{hours} hours, {mins} minutes and {secs} seconds"
        else:
            r["duration_fmt"] = f"{mins} minutes and {secs} seconds"
        r["created_fmt"]  = datetime.fromtimestamp(r["created"]/1000, tz=timezone.utc).strftime("%H:%M on %d/%m/%Y") if r.get("created") else "Unknown"
        r["created_time"] = datetime.fromtimestamp(r["created"]/1000, tz=timezone.utc).strftime("%H:%M") if r.get("created") else ""
        r["created_date"] = datetime.fromtimestamp(r["created"]/1000, tz=timezone.utc).strftime("%d/%m/%Y") if r.get("created") else ""
        return r
    except Exception as e:
        print(f"SQLite error: {e}"); return None

# ── Wasabi cloud storage ──────────────────────────────────────────────────────

def _wasabi_client():
    import boto3
    from botocore.client import Config
    return boto3.client(
        "s3",
        endpoint_url=f"https://s3.{WASABI_REGION}.wasabisys.com",
        aws_access_key_id=WASABI_ACCESS_KEY,
        aws_secret_access_key=WASABI_SECRET_KEY,
        config=Config(signature_version="s3v4"),
    )

def _normalise_slug(name):
    """Normalise bookmark name to match Nx Witness Wasabi folder naming.
    Strips [C]/[D] upload tags, lowercases, replaces non-alphanumeric with hyphens."""
    import re
    s = name.lower().strip()
    s = re.sub(r'\s*\[[cdCD]\]\s*', '', s)   # remove [C]/[D] tags (case-insensitive)
    s = re.sub(r'[^a-z0-9]+', '-', s)         # spaces/specials → hyphens
    s = re.sub(r'-+', '-', s).strip('-')       # collapse duplicates
    return s

def check_wasabi_footage(bookmark_slug):
    """Search Wasabi for a folder matching bookmark_slug containing video + JSON.
    Returns (True, prefix_str, arrived_str) or (False, None, None).
    arrived_str: HH:MM on DD/MM/YYYY — latest LastModified across video+JSON files,
    representing when the cloud package was complete and available.
    Prefix format: RMBC060/automated/DD-MM-YYYY/bookmark-slug/
    # TODO: cross-pipeline review — currently always checks RMBC060/automated/
    # for all pipelines. FOI and SYP may need separate bucket paths in future.
    """
    if not WASABI_ACCESS_KEY or not WASABI_SECRET_KEY:
        return False, None, None
    try:
        s3 = _wasabi_client()
        prefix = "RMBC060/automated/"
        paginator = s3.get_paginator("list_objects_v2")
        has_video       = False
        has_json        = False
        found_prefix    = None
        latest_modified = None   # tracks when the package became complete
        # Try exact slug first, then normalised version
        normalised = _normalise_slug(bookmark_slug)
        needles = list(dict.fromkeys([
            f"/{bookmark_slug}/",
            f"/{normalised}/",
        ]))
        for page in paginator.paginate(Bucket=WASABI_BUCKET, Prefix=prefix):
            for obj in page.get("Contents", []):
                key = obj["Key"]
                matched_needle = next((n for n in needles if n in key), None)
                if matched_needle:
                    if found_prefix is None:
                        idx = key.index(matched_needle)
                        found_prefix = key[: idx + len(matched_needle)]
                    ext = key.rsplit(".", 1)[-1].lower()
                    if ext in ("mkv", "mp4"):
                        has_video = True
                    elif ext == "json":
                        has_json = True
                    lm = obj.get("LastModified")
                    if lm and (latest_modified is None or lm > latest_modified):
                        latest_modified = lm
        if has_video and has_json and found_prefix:
            arrived = latest_modified.strftime("%H:%M on %d/%m/%Y") if latest_modified else ""
            return True, found_prefix, arrived
        # Partial match fallback — folder name contains normalised slug
        for page in paginator.paginate(Bucket=WASABI_BUCKET, Prefix=prefix):
            for obj in page.get("Contents", []):
                key = obj["Key"]
                parts = key.split("/")
                # parts: ['RMBC060','automated','DD-MM-YYYY','folder-name','file']
                if len(parts) >= 5 and normalised in parts[3]:
                    fp = "/".join(parts[:4]) + "/"
                    ext = key.rsplit(".", 1)[-1].lower()
                    if ext in ("mkv", "mp4"):
                        has_video = True
                        found_prefix = fp
                    elif ext == "json":
                        has_json = True
                        found_prefix = fp
                    lm = obj.get("LastModified")
                    if lm and (latest_modified is None or lm > latest_modified):
                        latest_modified = lm
        if has_video and has_json and found_prefix:
            arrived = latest_modified.strftime("%H:%M on %d/%m/%Y") if latest_modified else ""
            return True, found_prefix, arrived
        return False, None, None
    except Exception as e:
        print(f"Wasabi check error: {e}")
        return False, None, None

def upload_to_wasabi(docx_bytes, wasabi_prefix, filename):
    """Upload statement docx to Wasabi alongside its footage.
    Returns (timestamp_str, wasabi_key).
    """
    s3  = _wasabi_client()
    key = f"{wasabi_prefix}{filename}"
    s3.put_object(
        Bucket=WASABI_BUCKET,
        Key=key,
        Body=docx_bytes,
        ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    return datetime.now().strftime("%H:%M on %d/%m/%Y"), key

# ── Claude — Witness Statement ────────────────────────────────────────────────

def generate_statement(bm, form, download_time):

    if not OLLAMA_MODEL:
        raise ValueError(
            "OLLAMA_MODEL is not set. "
            "Add OLLAMA_MODEL=llama3.1 to /opt/CCTV_Statement/.env and restart the service."
        )

    # Reference line
    ref_parts = []
    if form.get("crime_name"):   ref_parts.append(form["crime_name"])
    if form.get("crime_number"): ref_parts.append(f"crime reference {form['crime_number']}")
    if form.get("flare_ref"):    ref_parts.append(f"Flare reference {form['flare_ref']}")
    if form.get("foi_ref"):      ref_parts.append(f"FOI reference {form['foi_ref']}")
    ref_line = " / ".join(ref_parts) if ref_parts else None

    # Bookmark creator
    bookmark_creator = form.get("bookmark_creator", "").strip()
    witness_name_val = form.get("witness_name", "").strip()
    creator_phrase = "I" if (not bookmark_creator or bookmark_creator == witness_name_val) else bookmark_creator

    # Clock narrative
    clock_date      = form.get("clock_check_date", download_time.strftime("%d/%m/%Y"))
    bt_time         = form.get("bt_clock_time", "")
    nvr_time        = form.get("nvr_clock_time", "")
    clock_diff      = form.get("clock_difference", "")
    clock_fast_slow = form.get("clock_fast_slow", "fast")

    if form.get("clock_checked") == "yes" and bt_time and nvr_time:
        if clock_fast_slow == "accurate":
            clock_conclusion = "showing the system time to be correct."
        elif clock_fast_slow == "fast":
            clock_conclusion = f"showing the system time to be approximately {clock_diff} fast."
        else:
            clock_conclusion = f"showing the system time to be approximately {clock_diff} slow."
        clock_para = (
            f"On {clock_date} I checked the NVR system clock against the BT Speaking Clock. "
            f"At {bt_time} on the BT Speaking Clock, the NVR system displayed {nvr_time}, "
            f"{clock_conclusion}"
        )
    else:
        clock_para = (
            "I was unable to verify the NVR system clock against the BT Speaking Clock at the time "
            "of this statement. The investigating officer should treat all timestamps as approximate "
            "and confirm NTP synchronisation status with the CCTV team."
        )

    # Export
    export_date = form.get("export_date", download_time.strftime("%d/%m/%Y"))
    export_time = form.get("export_time", download_time.strftime("%H:%M:%S"))
    exhibit_ref = form.get("exhibit_ref", "DCP1")
    media_type  = form.get("media_type", "Raw Data")
    is_dems     = media_type == "South Yorkshire Police DEMS Portal"
    secure_store = "Rawmarsh Police Station, Green Lane, Rawmarsh, Rotherham"

    export_para = (
        f"On {export_date} at {export_time} I exported the bookmarked footage from the Nx Witness system. "
        f"The footage was exported in MP4 format and saved to the following transferable media: {media_type}. "
        f"The footage was assigned exhibit reference {exhibit_ref}."
    )

    # Handover — all pipeline types
    handover_type = form.get("handover_type", "storage")

    if handover_type == "person":
        officer_name   = form.get("officer_name", "").strip() or "______________________"
        officer_number = form.get("officer_number", "").strip() or "______________________"
        h_location     = form.get("handover_location", "").strip() or "______________________"
        h_date         = form.get("handover_date", "").strip() or "______________________"
        h_time         = form.get("handover_time", "").strip() or "______________________"
        h_time_fmt = h_time if h_time else "__ : __ Hrs"
        handover = (
            f"On {h_date} at {h_time_fmt}, exhibit {exhibit_ref} was handed to {officer_name}, "
            f"collar number {officer_number}, at {h_location}."
        )

    elif handover_type == "dems":
        eu_date  = form.get("electronic_date", "").strip() or "__ / __ / _____"
        eu_time  = form.get("electronic_time", "").strip() or "__ : __ Hrs"
        eu_by    = form.get("electronic_by",   "").strip() or form.get("witness_name", "______________________")
        eu_recip = form.get("electronic_recipient", "").strip() or "______________________"
        eu_ref   = form.get("electronic_ref",  "").strip() or "______________________"
        handover = (
            f"The footage was uploaded to the South Yorkshire Police Digital Evidence Management System (DEMS)\n"
            f"- Date of upload: {eu_date}\n"
            f"- Time of upload: {eu_time}\n"
            f"- Uploaded by: {eu_by}\n"
            f"- Requester Details: {eu_recip}\n"
            f"- DEMS reference: {eu_ref}"
        )

    elif handover_type == "casefile":
        cf_ref    = form.get("casefile_ref", "").strip() or "______________________"
        cf_system = form.get("casefile_system", "Flare").strip()
        handover = (
            f"Exhibit {exhibit_ref} has been stored on the {cf_system}. "
            f"Internal case file reference: {cf_ref}."
        )

    else:
        # Secure storage locker
        if form.get("_pipeline") == "rmbc":
            # RMBC is internal-only — no police collection details
            handover = (
                f"Exhibit {exhibit_ref} has been placed into the secure CCTV storage locker at {secure_store}, "
                f"pending internal retrieval."
            )
        else:
            # SYP — police collection details required
            handover = (
                f"Exhibit {exhibit_ref} has been placed into the secure CCTV storage locker at {secure_store}, "
                f"awaiting collection. Handover details to be completed on collection:\n"
                f"- Officer name: ______________________\n"
                f"- Officers station: ______________________\n"
                f"- Collar / warrant number: ______________________\n"
                f"- Date and time of collection: __ / __ / _____ at __ : __ Hrs\n"
                f"- Location of handover: ______________________"
            )

    prompt = f"""You are a transcription tool. Your only job is to assemble the sections below into a flowing MG11 witness statement. You do NOT rewrite, rephrase, improve, or add to any of the provided text.

ABSOLUTE RULES — NO EXCEPTIONS:
- Copy the VERBATIM sections below WORD FOR WORD, CHARACTER FOR CHARACTER — do not change a single word
- Only Para 1 (introduction) is written by you — keep it to 2 sentences maximum
- Maximum 7 paragraphs total, no headings, no bold, no bullet points except where they appear in the verbatim text
- First person, past tense
- End with the exact Section 9 CJA 1967 declaration text: "This statement is true to the best of my knowledge and belief and I make it knowing that, if it is tendered in evidence, I shall be liable to prosecution if I have wilfully stated in it anything which I know to be false or do not believe to be true."
- Output plain text only — no markdown, no formatting

STRUCTURE:
Para 1 — Who I am, my role, where I am currently based, purpose of this statement{f' in relation to: {ref_line}' if ref_line else ''} (2 sentences max — write this yourself)
Para 2 — COPY THE SYSTEM DESCRIPTION VERBATIM BELOW — every word exactly as written
Para 3 — COPY THE BOOKMARK CREATION TEXT VERBATIM BELOW — every word exactly as written
Para 4 — COPY THE CLOCK CHECK TEXT VERBATIM BELOW — every word exactly as written
Para 5 — COPY THE EXPORT & HANDOVER TEXT VERBATIM BELOW — every word and every bullet line exactly as written
Para 6 — Section 9 CJA 1967 declaration (exact wording above)

=== WITNESS ===
Name: {form.get('witness_name')}
Role: {form.get('witness_role')}
Organisation: Rotherham Metropolitan Borough Council
Currently based at: {form.get('witness_base')}
Date of statement: {form.get('statement_date')}

=== INCIDENT ===
Start: {bm.get('start_fmt')}
End: {bm.get('end_fmt')}
Duration: {bm.get('duration_fmt')}
Bookmark name: {bm.get('name')}
Bookmark created: {bm.get('created_fmt')}

=== SYSTEM DESCRIPTION — COPY THIS VERBATIM, DO NOT CHANGE ANY WORD ===
The CCTV system at this location is referenced as {SITE_REF}. The system is owned and operated by Rotherham Metropolitan Borough Council.
The footage referred to in this statement was captured by a camera installed on {form.get('camera_location', 'Lamp Post')} at {form.get('incident_location', 'the above location')}.

=== BOOKMARK CREATION — COPY THIS VERBATIM, DO NOT CHANGE ANY WORD ===
{creator_phrase} created a bookmark within the system named "{bm.get('name')}" at {bm.get('created_time')} on {bm.get('created_date')} to preserve the relevant footage for export.

=== CLOCK CHECK — COPY THIS VERBATIM, DO NOT CHANGE ANY WORD ===
{clock_para}

=== EXPORT & HANDOVER — COPY THIS VERBATIM INCLUDING ALL BULLET LINES, DO NOT CHANGE ANY WORD ===
{export_para}

{handover}"""

    system_prompt = (
        "You are a precise transcription tool. You copy verbatim sections exactly as given — "
        "not one word changed. You only write the brief introduction paragraph yourself. "
        "No creativity. No improvements. No rephrasing. Plain text output only."
    )

    payload = {
        "model": OLLAMA_MODEL,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": prompt},
        ],
        "stream": False,
        "options": {"temperature": 0},
    }
    r = requests.post(f"{OLLAMA_HOST}/api/chat", json=payload, timeout=300)
    r.raise_for_status()
    return r.json()["message"]["content"]

# ── FOI Audit Document — data assembler (no AI — deterministic legal doc) ──────

def assemble_foi_data(bm, form, download_time):
    """Assembles all FOI disclosure record fields into a structured dict."""

    request_type      = form.get("foi_request_type", "Public")
    identifiable = form.get("foi_identifiable", "no") == "yes"

    # Legal basis logic
    if identifiable:
        legal_basis = (
            "Data Protection — UK GDPR Article 6(1)(f) (legitimate interests) and/or "
            "Article 6(1)(c) (legal obligation). DPA 2018 Schedule 2 applies where "
            "disclosure is required for legal proceedings or law enforcement purposes. "
            "Note: This disclosure involves footage containing identifiable individuals "
            "and must NOT be processed under FOIA 2000."
        )
        governing_legislation = "Data Protection Act 2018 / UK GDPR"
    elif request_type == "Public":
        legal_basis = (
            "Freedom of Information Act 2000. Disclosure is made in line with the Freedom of "
            "Information Act 2000 and UK GDPR security and accountability principles (Article 5). "
            "No identifiable individuals present in disclosed footage. Where personal data may be "
            "incidentally captured, Section 40 FOIA exemption has been considered."
        )
        governing_legislation = "Freedom of Information Act 2000 / Data Protection Act 2018"
    else:  # Solicitor or Insurance
        legal_basis = (
            "Data Protection Act 2018 Schedule 2 (legal claims). UK GDPR Article 6(1)(f) "
            "(legitimate interests of the requesting party in connection with legal proceedings). "
            "Disclosure is made in line with UK GDPR security and accountability principles (Article 5). "
            "This disclosure is not made under FOIA 2000."
        )
        governing_legislation = "Data Protection Act 2018 / UK GDPR"

    # Redaction — engineer's opinion only, FOI team decides
    redaction_onsite = form.get("redaction_onsite", "yes")
    if redaction_onsite == "no":
        redaction_further      = "No — engineer assessment: footage appears suitable for disclosure"
        redaction_responsibility = "FOI Team / Information Governance to confirm before onward disclosure."
    else:
        redaction_further      = "Yes — engineer assessment: footage likely requires redaction"
        redaction_responsibility = "FOI Team / Information Governance — footage must be reviewed before onward disclosure."

    # Time verification
    time_verified = form.get("time_verified", "no") == "yes"

    return {
        "ref_id":               form.get("foi_ref", "Not provided"),
        "date_disclosure":      download_time.strftime("%d/%m/%Y"),
        "time_disclosure":      download_time.strftime("%H:%M"),
        "request_type":         request_type,
        "requestor":            form.get("foi_requester", "Not provided"),
        "organisation":         form.get("foi_organisation", "Not provided"),
        "date_received":        form.get("foi_date_received", "Not provided"),
        "request_purpose":      form.get("foi_summary", "Not provided"),
        "governing_legislation":governing_legislation,
        "legal_basis":          legal_basis,
        "identifiable":         "Yes" if identifiable else "No",
        "location":             form.get("incident_location", "Not provided"),
        "incident_type":        form.get("foi_incident_type", "Not provided"),
        "system_ref":           SITE_REF,
        "footage_start":        bm.get("start_fmt", ""),
        "footage_end":          bm.get("end_fmt", ""),
        "footage_duration":     bm.get("duration_fmt", ""),
        "export_format":        form.get("foi_export_format", "MP4"),
        "delivery_method":      form.get("media_type", "Not specified"),
        "encryption":           form.get("foi_encryption", "No"),
        "viewing_software":     form.get("foi_viewing_software", "No"),
        "exhibit_id":           form.get("exhibit_ref", "Not provided"),
        "redaction_onsite":     "Yes" if redaction_onsite == "yes" else "No",
        "redaction_further":    redaction_further,
        "redaction_resp":       redaction_responsibility,
        "time_verified":        time_verified,
        "time_source":          "BT Speaking Clock",
        "time_verification":    form.get("foi_verify_time", "") if time_verified else "",
        "time_system":          form.get("foi_system_time", "") if time_verified else "",
        "time_offset":          form.get("foi_time_offset", "") if time_verified else "",
        "verified_by":          form.get("witness_name", "") if time_verified else "",
        "disclosed_by":         form.get("witness_name", ""),
        "disclosed_role":       form.get("witness_role", ""),
        "statement_date":       form.get("statement_date", download_time.strftime("%d/%m/%Y")),
    }

# ── Word doc — Witness Statement ──────────────────────────────────────────────

def build_docx(statement_text, form, download_time):
    witness_name = form.get("witness_name", "Officer")
    doc = Document()

    # Page setup
    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)

    # Header
    hp = doc.sections[0].header.paragraphs[0]
    hp.text = "OFFICIAL SENSITIVE  —  WITNESS STATEMENT  —  NOT FOR DISTRIBUTION"
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.runs[0].font.size = Pt(8)
    hp.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    hp.runs[0].font.name = "Calibri"

    # Footer
    fp = doc.sections[0].footer.paragraphs[0]
    fp.text = f"Generated: {download_time.strftime('%d/%m/%Y %H:%M')}  |  RMBC CCTV Evidence Unit  |  Prepared by: {witness_name}  |  System: {SITE_REF}"
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.runs[0].font.size = Pt(8)
    fp.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    fp.runs[0].font.name = "Calibri"

    def add_rule(doc, color=(0xCC, 0xCC, 0xCC)):
        rule = doc.add_paragraph()
        rule.paragraph_format.space_before = Pt(2)
        rule.paragraph_format.space_after  = Pt(4)
        rr = rule.add_run("─" * 90)
        rr.font.size = Pt(7)
        rr.font.color.rgb = RGBColor(*color)

    def add_para(doc, text, size=11, space_after=8, bold=False, italic=False, color=None, align=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(2)
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = "Calibri"
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)
        if align:
            p.alignment = align
        return p

    # ── Title block ──────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(6)
    t.paragraph_format.space_after  = Pt(4)
    tr = t.add_run("WITNESS STATEMENT")
    tr.bold = True
    tr.font.size = Pt(20)
    tr.font.name = "Calibri"
    tr.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(2)
    sr = sub.add_run("CJ Act 1967, s.9  ·  MC Act 1980, ss.5A(3)(a) and 5B  ·  Police and Criminal Evidence Act 1984")
    sr.italic = True
    sr.font.size = Pt(9)
    sr.font.name = "Calibri"
    sr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    add_rule(doc, color=(0x1F, 0x49, 0x7D))

    # ── Reference metadata block ─────────────────────────────────────────────
    ref_parts = []
    if form.get("crime_name"):   ref_parts.append(f"Incident: {form['crime_name']}")
    if form.get("crime_number"): ref_parts.append(f"Crime Ref: {form['crime_number']}")
    if form.get("flare_ref"):    ref_parts.append(f"Flare Ref: {form['flare_ref']}")
    if form.get("foi_ref"):      ref_parts.append(f"FOI Ref: {form['foi_ref']}")
    if ref_parts:
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.paragraph_format.space_after = Pt(2)
        mr = meta.add_run("  |  ".join(ref_parts))
        mr.font.size = Pt(9)
        mr.font.name = "Calibri"
        mr.bold = True
        mr.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    stmt_date = form.get("statement_date", download_time.strftime("%d/%m/%Y"))
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_after = Pt(8)
    dr = date_p.add_run(f"Statement Date: {stmt_date}  |  System Reference: {SITE_REF}")
    dr.font.size = Pt(9)
    dr.font.name = "Calibri"
    dr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    add_rule(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Statement body ───────────────────────────────────────────────────────
    for line in statement_text.split("\n"):
        line = line.rstrip()
        if not line or line.startswith("---"):
            # Minimal gap between paragraphs — not double spacing
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(line[2:])
            run.font.size = Pt(11)
            run.font.name = "Calibri"
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(10)
            p.paragraph_format.space_before = Pt(0)
            parts = line.split("**")
            bold = False
            for part in parts:
                if part:
                    run = p.add_run(part)
                    run.bold = bold
                    run.font.size = Pt(11)
                    run.font.name = "Calibri"
                bold = not bold

    # ── Divider before signature ─────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    add_rule(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── Signature block ──────────────────────────────────────────────────────
    def sig_line(doc, label, value="", line_len=20):
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(12)
        p.paragraph_format.space_before = Pt(0)
        lr = p.add_run(f"{label}  ")
        lr.bold = True; lr.font.size = Pt(11); lr.font.name = "Calibri"
        vr = p.add_run(value if value else "_" * line_len)
        vr.font.size = Pt(11); vr.font.name = "Calibri"

    sig_line(doc, "Signed:", line_len=20)
    sig_line(doc, "Date:", value=stmt_date)
    sig_line(doc, "Print Name:", value=witness_name.upper())

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ── Word doc — FOI Disclosure Record ─────────────────────────────────────────

def build_foi_docx(data, download_time):
    """Builds the 12-section CCTV Disclosure Record Word document."""
    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Inches(0.9)
        sec.bottom_margin = Inches(0.9)
        sec.left_margin   = Inches(1.2)
        sec.right_margin  = Inches(1.2)

    # Header
    hp = doc.sections[0].header.paragraphs[0]
    hp.text = "OFFICIAL SENSITIVE  ·  CCTV DISCLOSURE RECORD  ·  INTERNAL USE ONLY"
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.runs[0].font.size = Pt(8)
    hp.runs[0].font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
    hp.runs[0].font.name = "Calibri"

    # Footer
    fp = doc.sections[0].footer.paragraphs[0]
    fp.text = (f"Generated: {download_time.strftime('%d/%m/%Y %H:%M')}  ·  "
               f"RMBC CCTV Evidence Unit  ·  {data['disclosed_by']}  ·  {SITE_REF}")
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.runs[0].font.size = Pt(8)
    fp.runs[0].font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
    fp.runs[0].font.name = "Calibri"

    BLUE    = RGBColor(0x1F, 0x49, 0x7D)
    MIDBLUE = RGBColor(0x2E, 0x6B, 0xB0)
    GREY    = RGBColor(0x55, 0x55, 0x55)
    LGREY   = RGBColor(0xCC, 0xCC, 0xCC)

    from docx.oxml.ns import qn as _qn
    from docx.oxml   import OxmlElement as _OXE

    def page_break():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run()
        br = _OXE("w:br")
        br.set(_qn("w:type"), "page")
        run._r.append(br)

    def rule(color=LGREY, thick=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(5)
        r = p.add_run("─" * 95)
        r.font.size = Pt(7) if not thick else Pt(8)
        r.font.color.rgb = color

    def section_heading(num, title, pb=False):
        if pb:
            page_break()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(f"{num}.  {title.upper()}")
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.name = "Calibri"
        r.font.color.rgb = BLUE
        rule(color=MIDBLUE)

    def field(label, value, indent=True):
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.space_before = Pt(0)
        if indent:
            p.paragraph_format.left_indent = Inches(0.25)
        lr = p.add_run(f"{label}:  ")
        lr.bold = True
        lr.font.size = Pt(10)
        lr.font.name = "Calibri"
        lr.font.color.rgb = GREY
        vr = p.add_run(str(value) if value else "\u2014")
        vr.font.size = Pt(10)
        vr.font.name = "Calibri"

    def note(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.25)
        r = p.add_run(text)
        r.italic = True
        r.font.size = Pt(9.5)
        r.font.name = "Calibri"
        r.font.color.rgb = GREY

    # ── Title block
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(10)
    t.paragraph_format.space_after  = Pt(3)
    tr = t.add_run("CCTV DISCLOSURE RECORD")
    tr.bold = True; tr.font.size = Pt(22)
    tr.font.name = "Calibri"; tr.font.color.rgb = BLUE

    cl = doc.add_paragraph()
    cl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cl.paragraph_format.space_after = Pt(1)
    cr = cl.add_run("Classification: OFFICIAL SENSITIVE  ·  Disclosure Audit & Evidential Record")
    cr.italic = True; cr.font.size = Pt(9)
    cr.font.name = "Calibri"; cr.font.color.rgb = GREY

    leg = doc.add_paragraph()
    leg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    leg.paragraph_format.space_after = Pt(6)
    lr2 = leg.add_run(f"Governing Legislation: {data['governing_legislation']}")
    lr2.italic = True; lr2.font.size = Pt(9)
    lr2.font.name = "Calibri"; lr2.font.color.rgb = GREY

    rule(color=BLUE, thick=True)

    # S1
    section_heading("1", "Disclosure Reference")
    field("Reference ID", data["ref_id"])
    field("Date of Disclosure", f"{data['date_disclosure']} at {data['time_disclosure']}")

    # S2
    section_heading("2", "Request Details")
    field("Request Type", data["request_type"])
    field("Requestor Name", data["requestor"])
    field("Requesting Organisation", data["organisation"])
    field("Date Request Received", data["date_received"])
    field("Request Purpose", data["request_purpose"])

    # S3
    section_heading("3", "Legal Basis for Disclosure")
    field("Identifiable Individuals or Vehicles in Footage", data["identifiable"])
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent  = Inches(0.25)
    r = p.add_run(data["legal_basis"])
    r.font.size = Pt(9.5); r.font.name = "Calibri"; r.font.color.rgb = GREY

    # S4
    section_heading("4", "Incident Details")
    field("Location", data["location"])
    field("Incident Type", data["incident_type"])
    field("System Reference", data["system_ref"])

    # S5
    section_heading("5", "Footage Details")
    field("Start Time", data["footage_start"])
    field("End Time", data["footage_end"])
    field("Duration", data["footage_duration"])

    # S6 — page break before
    section_heading("6", "Disclosure Scope Statement", pb=True)
    note("Only footage directly relevant to the stated request has been disclosed. "
         "No additional footage has been provided beyond the period and location specified above. "
         "This disclosure does not constitute admission of liability on the part of "
         "Rotherham Metropolitan Borough Council.")

    # S7
    section_heading("7", "Format & Method of Disclosure")
    field("Export Format", data["export_format"])
    field("Delivery Method", data["delivery_method"])
    field("Encryption Applied", data["encryption"])
    field("Viewing Software Provided", data["viewing_software"])

    # S8
    section_heading("8", "Exhibit Reference")
    field("Exhibit ID", data["exhibit_id"])

    # S9
    section_heading("9", "Handling & Integrity Statement")
    note("Footage was extracted directly from the Nx Witness video management system. "
         "No alterations have been made to the footage unless redaction is explicitly stated below. "
         "Evidence continuity has been maintained throughout the handling process. "
         "No hashing or digital signing has been applied unless separately documented.")

    # S10
    section_heading("10", "Redaction & Data Protection Handling")
    field("Does Footage Require Redacting", data["redaction_onsite"])
    field("Redaction Responsibility", data["redaction_resp"])

    # S11
    section_heading("11", "Time Synchronisation Verification")
    if data["time_verified"]:
        field("Time Verification Performed", "Yes")
        field("Verification Source", data["time_source"])
        field("Verification Time (Reference Clock)", data["time_verification"])
        field("System Time at Verification", data["time_system"])
        field("Offset", data["time_offset"])
        field("Verified By", data["verified_by"])
    else:
        field("Time Verification Performed", "No")
        note("Time verification was not performed for this disclosure. "
             "Timestamps should be treated as system-reported and may vary from "
             "absolute time if NTP synchronisation has not been confirmed.")

    # S12 — page break before, on its own final page
    section_heading("12", "Authorisation", pb=True)
    field("Disclosed By", f"{data['disclosed_by']} — {data['disclosed_role']}")
    field("Organisation", "Rotherham Metropolitan Borough Council — CCTV Evidence Unit")

    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    rule(color=BLUE, thick=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    for label, val in [
        ("Signature", ""),
        ("Print Name", data["disclosed_by"].upper()),
        ("Date", data["statement_date"]),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(16)
        lr = p.add_run(f"{label}:  ")
        lr.bold = True; lr.font.size = Pt(11); lr.font.name = "Calibri"
        lr.font.color.rgb = GREY
        vr = p.add_run(val if val else "_" * 45)
        vr.font.size = Pt(11); vr.font.name = "Calibri"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ── Templates ─────────────────────────────────────────────────────────────────

LOGIN_HTML = """<!DOCTYPE html>
<html><head><title>RMBC CCTV — Sign In</title>
<meta name="viewport" content="width=device-width, initial-scale=1">
<link href="https://fonts.googleapis.com/css2?family=DM+Mono:wght@400;500&family=DM+Sans:wght@400;500;700&display=swap" rel="stylesheet">
<style>
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:'DM Sans',sans-serif;background:#0d1117;min-height:100vh;display:flex;align-items:center;justify-content:center;}
.card{background:#161b22;border:1px solid #30363d;border-radius:12px;padding:40px;width:420px;}
.logo{text-align:center;margin-bottom:32px;}
.logo .icon{font-size:40px;margin-bottom:12px;}
.logo h1{color:#e6edf3;font-size:22px;font-weight:700;}
.logo p{color:#8b949e;font-size:13px;margin-top:4px;}
.badge{background:#1f4068;color:#58a6ff;padding:3px 10px;border-radius:20px;font-size:11px;font-family:'DM Mono',monospace;display:inline-block;margin-bottom:16px;}
label{display:block;font-size:12px;color:#8b949e;font-weight:500;margin-bottom:6px;margin-top:16px;text-transform:uppercase;letter-spacing:0.5px;}
input{width:100%;padding:10px 14px;background:#0d1117;border:1px solid #30363d;border-radius:6px;font-size:14px;color:#e6edf3;font-family:'DM Sans',sans-serif;}
input:focus{outline:none;border-color:#58a6ff;}
button{width:100%;padding:12px;background:#238636;color:white;border:none;border-radius:6px;font-size:14px;font-weight:600;cursor:pointer;margin-top:20px;font-family:'DM Sans',sans-serif;}
button:hover{background:#2ea043;}
.error{background:#3d1a1a;border:1px solid #f85149;color:#f85149;padding:10px 14px;border-radius:6px;margin-bottom:16px;font-size:13px;}
.footer{text-align:center;margin-top:20px;font-size:11px;color:#484f58;}
</style></head>
<body><div class="card">
<div class="logo">
    <div class="icon">🎥</div>
    <span class="badge">OFFICIAL SENSITIVE</span>
    <h1>RMBC CCTV</h1>
    <p>Evidence Management System</p>
</div>
{% if error %}<div class="error">{{ error }}</div>{% endif %}
<form method="POST">
    <label>Username</label>
    <input type="text" name="username" placeholder="e.g. dane.plant" autofocus required>
    <label>Password</label>
    <input type="password" name="password" required>
    <button type="submit">Sign In →</button>
</form>
<div class="footer">Rotherham Metropolitan Borough Council · CCTV Evidence Unit · {{ site_ref }}</div>
</div></body></html>"""

BOOKMARKS_HTML = """<!DOCTYPE html>
<html><head><title>RMBC CCTV — Select Incident</title>
<meta name="viewport" content="width=device-width, initial-scale=1">
<link href="https://fonts.googleapis.com/css2?family=DM+Mono:wght@400;500&family=DM+Sans:wght@400;500;700&display=swap" rel="stylesheet">
<style>
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:'DM Sans',sans-serif;background:#0d1117;min-height:100vh;color:#e6edf3;}
.topbar{background:#161b22;border-bottom:1px solid #30363d;padding:14px 20px;display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px;}
.topbar h1{font-size:15px;font-weight:700;}
.topbar .right{font-size:13px;color:#8b949e;}
.topbar a{color:#58a6ff;text-decoration:none;margin-left:12px;font-size:13px;}
.wordmark{font-size:10px;font-weight:700;letter-spacing:3px;color:#58a6ff;font-family:'DM Mono',monospace;text-transform:uppercase;line-height:1;margin-bottom:3px;}
.site-badge{background:#1f2937;color:#58a6ff;padding:2px 8px;border-radius:4px;font-size:11px;font-family:'DM Mono',monospace;margin-left:8px;}
.container{max-width:900px;margin:24px auto;padding:0 16px;}
.page-title{font-size:20px;font-weight:700;margin-bottom:4px;}
.page-sub{color:#8b949e;font-size:13px;margin-bottom:20px;}
.bm{background:#161b22;border:1px solid #30363d;border-radius:10px;padding:16px 20px;margin-bottom:10px;display:flex;justify-content:space-between;align-items:center;border-left:4px solid #30363d;gap:12px;flex-wrap:wrap;}
.bm:hover{border-color:#58a6ff;border-left-color:#58a6ff;}
.bm-cloud{background:#081910;border-color:#2ea043;border-left:5px solid #3fb950;}
.bm-cloud .bm-name{color:#cae8ca;}
.bm-cloud:hover{border-color:#3fb950;}
.bm-pending{background:#191000;border-color:#9e6a03;border-left:5px solid #d29922;}
.bm-pending .bm-name{color:#e8c87a;}
.bm-pending:hover{border-color:#e3b341;}
.bm-info{flex:1;min-width:0;}
.bm-name{font-size:15px;font-weight:600;margin-bottom:4px;}
.bm-desc{font-size:13px;color:#8b949e;margin-bottom:4px;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;}
.bm-time{font-size:11px;color:#484f58;font-family:'DM Mono',monospace;}
.bm-status{display:inline-flex;align-items:center;gap:5px;padding:3px 9px;border-radius:4px;font-size:11px;font-weight:700;font-family:'DM Mono',monospace;margin-bottom:6px;}
.bm-status-confirmed{background:#0d2b0d;color:#3fb950;border:1px solid #238636;}
.bm-status-pending{background:#2d1f00;color:#d29922;border:1px solid #9e6a03;}
.bm-status-none{background:#1c2128;color:#484f58;border:1px solid #30363d;}
.bm-right{display:flex;flex-direction:column;align-items:flex-end;gap:6px;flex-shrink:0;}
.dur{font-size:13px;font-weight:700;color:#58a6ff;font-family:'DM Mono',monospace;}
.btn-row{display:flex;gap:6px;flex-wrap:wrap;justify-content:flex-end;}
.btn{padding:6px 14px;border-radius:6px;font-size:12px;font-weight:600;text-decoration:none;display:inline-block;white-space:nowrap;}
.btn-syp{background:#238636;color:white;}
.btn-syp:hover{background:#2ea043;}
.btn-rmbc{background:#1f4068;color:#58a6ff;border:1px solid #30363d;}
.btn-rmbc:hover{background:#2a5a8a;color:white;}
.btn-foi{background:#2a1a4a;color:#8a5cf6;border:1px solid #6e40c9;}
.btn-foi:hover{background:#6e40c9;color:white;}
.id-badge{background:#1f2937;color:#484f58;padding:2px 6px;border-radius:4px;font-size:10px;font-family:'DM Mono',monospace;margin-left:6px;}
.cloud-tag{display:inline-flex;align-items:center;gap:3px;padding:1px 7px;border-radius:4px;font-size:10px;font-weight:700;font-family:'DM Mono',monospace;margin-left:6px;vertical-align:middle;}
.cloud-tag-c{background:#0d2b0d;color:#3fb950;border:1px solid #238636;}
.cloud-tag-d{background:#2d1f00;color:#d29922;border:1px solid #d29922;}
.empty{text-align:center;padding:60px;color:#484f58;}
@media(max-width:600px){
  .bm{flex-direction:column;align-items:flex-start;}
  .bm-right{align-items:flex-start;width:100%;}
  .btn-row{width:100%;}
  .btn{flex:1;text-align:center;}
  .sidebar{display:none;}
}
.sidebar-wrap{position:fixed;right:0;top:50vh;transform:translateY(-50%);z-index:9999;display:flex;align-items:stretch;}
.sidebar-drawer{background:#161b22;border:1px solid #30363d;border-right:none;border-radius:10px 0 0 10px;width:0;overflow:hidden;opacity:0;transition:width 0.25s ease,opacity 0.2s ease;display:flex;flex-direction:column;justify-content:flex-start;}
.sidebar-drawer.open{width:220px;opacity:1;}
.sidebar-inner{padding:16px;white-space:nowrap;min-width:220px;}
.sidebar-title{font-size:10px;color:#484f58;text-transform:uppercase;letter-spacing:1.5px;margin-bottom:10px;font-family:'DM Mono',monospace;}
.sidebar-link{display:block;font-size:12px;color:#58a6ff;text-decoration:none;padding:7px 10px;border-radius:6px;margin-bottom:4px;line-height:1.4;transition:background 0.15s;}
.sidebar-link:hover{background:#1f4068;}
.sidebar-divider{border:none;border-top:1px solid #21262d;margin:8px 0;}
.sidebar-tab{background:#1f2937;color:#58a6ff;border:1px solid #30363d;border-right:none;border-radius:10px 0 0 10px;padding:12px 7px;cursor:pointer;display:flex;align-items:center;justify-content:center;width:28px;flex-shrink:0;transition:background 0.2s;}
.sidebar-tab:hover{background:#1f4068;}
.sidebar-tab span{writing-mode:vertical-rl;text-orientation:mixed;font-size:10px;font-weight:700;letter-spacing:2px;text-transform:uppercase;font-family:'DM Sans',sans-serif;user-select:none;}
@media(max-width:768px){
  .sidebar-wrap{top:auto;bottom:16px;transform:none;right:16px;flex-direction:column-reverse;align-items:flex-end;}
  .sidebar-drawer{border-radius:10px;border-right:1px solid #30363d;width:0;}
  .sidebar-drawer.open{width:220px;margin-bottom:8px;}
  .sidebar-tab{border-radius:50px;border-right:1px solid #30363d;width:auto;padding:8px 14px;writing-mode:horizontal-tb;}
  .sidebar-tab span{writing-mode:horizontal-tb;letter-spacing:1px;font-size:11px;}
}
</style></head>
<body>

<div class="sidebar-wrap" id="sidebarWrap">
    <div class="sidebar-drawer" id="sidebarDrawer">
        <div class="sidebar-inner">
            <div class="sidebar-title">Resources</div>
            <a class="sidebar-link" href="https://ico.org.uk/for-organisations/uk-gdpr-guidance-and-resources/" target="_blank" rel="noopener">Current UK GDPR Guidance</a>
            <a class="sidebar-link" href="https://ico.org.uk/about-the-ico/" target="_blank" rel="noopener">About the ICO</a>
            <hr class="sidebar-divider">
            <a class="sidebar-link" href="https://www.legislation.gov.uk/ukpga/2000/36/contents" target="_blank" rel="noopener">Freedom of Information Act 2000</a>
            <a class="sidebar-link" href="https://www.legislation.gov.uk/ukpga/2018/12/contents" target="_blank" rel="noopener">Data Protection Act 2018</a>
            <hr class="sidebar-divider">
            <a class="sidebar-link" href="https://www.rotherham.gov.uk/consultation-feedback/freedom-information-request-foi" target="_blank" rel="noopener">RMBC FOI Policy</a>
        </div>
    </div>
    <div class="sidebar-tab" onclick="toggleSidebar()"><span>Resources</span></div>
</div>
<script>
function toggleSidebar(){
    var d=document.getElementById('sidebarDrawer');
    if(d) d.classList.toggle('open');
}
</script>
<div class="topbar">
    <div>
        <div class="wordmark">CamScribe</div>
        <h1>🎥 RMBC CCTV — Evidence Management<span class="site-badge">{{ site_ref }}</span></h1>
    </div>
    <div class="right">{{ session.user_name }} · {{ session.user_role }}<a href="/logout">Sign out</a></div>
</div>
<div class="container">
    <div class="page-title">Select Incident Bookmark</div>
    <div class="page-sub">Choose a bookmark to generate a Witness Statement or FOI Disclosure Record.</div>
    {% if bookmarks %}
        {% for bm in bookmarks %}
        <div class="bm{% if bm.wasabi_confirmed %} bm-cloud{% elif '[C]' in ((bm.name or '') + ' ' + (bm.description or '')).upper() or '[D]' in ((bm.name or '') + ' ' + (bm.description or '')).upper() %} bm-pending{% endif %}">
            <div class="bm-info">
                {% if bm.wasabi_confirmed %}
                <div class="bm-status bm-status-confirmed">&#10003; Cloud Enabled</div>
                {% elif '[C]' in ((bm.name or '') + ' ' + (bm.description or '')).upper() %}
                <div class="bm-status bm-status-pending">&#9729; Cloud bookmark — not yet in Wasabi</div>
                {% elif '[D]' in ((bm.name or '') + ' ' + (bm.description or '')).upper() %}
                <div class="bm-status bm-status-pending">&#8987; Deferred — scheduled for overnight upload</div>
                {% else %}
                <div class="bm-status bm-status-none">&#11015; No cloud tag — download only</div>
                {% endif %}
                <div class="bm-name">{{ bm.name or "(No name)" }}</div>
                <div class="bm-desc">{{ bm.description or "No description" }}</div>
                <div class="bm-time">{{ bm.start_fmt }} → {{ bm.end_fmt }}</div>
            </div>
            <div class="bm-right">
                <div class="dur">{{ bm.duration_fmt }}</div>
                <div class="btn-row">
                    <div style="display:flex;flex-direction:column;align-items:center;gap:4px;">
                        <a href="/syp/{{ bm.record_id }}" class="btn btn-syp">SYP Statement</a>
                        <span style="font-size:10px;color:#484f58;">Police investigation — includes DEMS</span>
                    </div>
                    <div style="display:flex;flex-direction:column;align-items:center;gap:4px;">
                        <a href="/rmbc/{{ bm.record_id }}" class="btn btn-rmbc">RMBC Statement</a>
                        <span style="font-size:10px;color:#484f58;">Internal use</span>
                    </div>
                    <div style="display:flex;flex-direction:column;align-items:center;gap:4px;">
                        <a href="/foi/{{ bm.record_id }}" class="btn btn-foi">FOI Record</a>
                        <span style="font-size:10px;color:#484f58;">Public, solicitor, or insurance disclosure</span>
                    </div>
                </div>
            </div>
        </div>
        {% endfor %}
    {% else %}
        <div class="empty">No bookmarks found. Create one in Nx Witness first.</div>
    {% endif %}
</div></body></html>"""

FORM_HTML = """<!DOCTYPE html>
<html><head><title>RMBC CCTV — Statement Form</title>
<meta name="viewport" content="width=device-width, initial-scale=1">
<link href="https://fonts.googleapis.com/css2?family=DM+Mono:wght@400;500&family=DM+Sans:wght@400;500;700&display=swap" rel="stylesheet">
<style>
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:'DM Sans',sans-serif;background:#0d1117;color:#e6edf3;}
.topbar{background:#161b22;border-bottom:1px solid #30363d;padding:14px 30px;display:flex;justify-content:space-between;align-items:center;}
.topbar h1{font-size:16px;font-weight:700;}
.topbar a{color:#58a6ff;text-decoration:none;font-size:13px;margin-left:16px;}
.wordmark{font-size:10px;font-weight:700;letter-spacing:3px;color:#58a6ff;font-family:'DM Mono',monospace;text-transform:uppercase;line-height:1;margin-bottom:3px;}
.container{max-width:820px;margin:30px auto;padding:0 20px 80px;}
.clock-box{background:#0d1117;border:1px solid #238636;border-radius:10px;padding:20px 24px;margin-bottom:24px;text-align:center;}
.clock-label{font-size:11px;color:#484f58;text-transform:uppercase;letter-spacing:1px;margin-bottom:8px;font-family:'DM Mono',monospace;}
.clock-date{font-size:17px;font-weight:600;color:#8b949e;margin-bottom:6px;font-family:'DM Mono',monospace;}
.clock-time{font-size:48px;font-weight:700;color:#58a6ff;letter-spacing:3px;font-family:'DM Mono',monospace;line-height:1;}
.clock-utc{font-size:11px;color:#484f58;margin-top:6px;font-family:'DM Mono',monospace;}
.incident{background:#161b22;border:1px solid #30363d;border-radius:10px;padding:18px 22px;margin-bottom:20px;border-left:3px solid #58a6ff;}
.incident h3{font-size:13px;font-weight:700;color:#58a6ff;margin-bottom:8px;text-transform:uppercase;letter-spacing:0.5px;}
.incident p{font-size:13px;color:#8b949e;line-height:1.8;font-family:'DM Mono',monospace;}
.section{background:#161b22;border:1px solid #30363d;border-radius:10px;padding:24px;margin-bottom:16px;}
.section h3{font-size:12px;font-weight:700;color:#8b949e;text-transform:uppercase;letter-spacing:1px;margin-bottom:20px;padding-bottom:10px;border-bottom:1px solid #21262d;}
.green-section{background:#0a1f0a;border:1px solid #238636;border-radius:10px;padding:24px;margin-bottom:16px;}
.green-section h3{font-size:12px;font-weight:700;color:#3fb950;text-transform:uppercase;letter-spacing:1px;margin-bottom:8px;padding-bottom:10px;border-bottom:1px solid #1a3a1a;}
.warning-box{background:#2d1f00;border:1px solid #d29922;border-radius:8px;padding:12px 16px;margin-top:14px;font-size:13px;color:#d29922;line-height:1.6;}
.warning-box strong{display:block;margin-bottom:4px;}
.hint{font-size:13px;color:#8b949e;margin-bottom:16px;line-height:1.5;}
label{display:block;font-size:12px;color:#8b949e;font-weight:500;margin-bottom:6px;margin-top:14px;text-transform:uppercase;letter-spacing:0.4px;}
label span{font-weight:400;text-transform:none;color:#484f58;margin-left:4px;}
input,select,textarea{width:100%;padding:10px 14px;background:#0d1117;border:1px solid #30363d;border-radius:6px;font-size:14px;color:#e6edf3;font-family:'DM Sans',sans-serif;}
input:focus,select:focus,textarea:focus{outline:none;border-color:#58a6ff;}
select option{background:#161b22;}
textarea{resize:vertical;min-height:80px;}
.row{display:grid;grid-template-columns:1fr 1fr;gap:14px;}
.row3{display:grid;grid-template-columns:1fr 1fr 1fr;gap:14px;}
.toggle-row{display:flex;align-items:center;gap:12px;margin-top:14px;}
.toggle-row input[type=checkbox]{width:auto;margin:0;}
.toggle-row label{margin:0;text-transform:none;font-size:13px;color:#8b949e;letter-spacing:0;}
.handover-opts{display:grid;grid-template-columns:1fr 1fr 1fr;gap:10px;margin-top:10px;}
.h-opt{border:1px solid #30363d;border-radius:8px;padding:14px;text-align:center;cursor:pointer;background:#0d1117;}
.h-opt:hover{border-color:#58a6ff;}
.h-opt.active{border-color:#238636;background:#0d2b0d;}
.h-opt input[type=radio]{display:none;}
.h-opt .icon{font-size:24px;margin-bottom:6px;}
.h-opt .title{font-size:13px;font-weight:600;color:#e6edf3;}
.h-opt .sub{font-size:11px;color:#484f58;margin-top:2px;}
.submit-btn{width:100%;padding:16px;background:#238636;color:white;border:none;border-radius:8px;font-size:16px;font-weight:700;cursor:pointer;margin-top:10px;font-family:'DM Sans',sans-serif;}
.submit-btn:hover{background:#2ea043;}
.submit-btn:disabled{background:#21262d;color:#484f58;cursor:not-allowed;}
#loadingBox{display:none;background:#161b22;border:1px solid #30363d;border-radius:10px;padding:30px;text-align:center;margin-top:16px;}
#successBox{display:none;}
@media(max-width:640px){
  .row,.row3{grid-template-columns:1fr !important;}
  .clock-time{font-size:30px !important;}
  .handover-opts{grid-template-columns:1fr !important;}
  .container{padding:0 12px 60px;}
  .topbar h1{font-size:13px;}
  .sidebar-wrap{top:auto;bottom:16px;transform:none;right:16px;flex-direction:column-reverse;align-items:flex-end;}
  .sidebar-drawer{border-radius:10px;border-right:1px solid #30363d;width:0;}
  .sidebar-drawer.open{width:220px;margin-bottom:8px;}
  .sidebar-tab{border-radius:50px;border-right:1px solid #30363d;width:auto;padding:8px 14px;}
  .sidebar-tab span{writing-mode:horizontal-tb;letter-spacing:1px;font-size:11px;}
}
.sidebar-wrap{position:fixed;right:0;top:50vh;transform:translateY(-50%);z-index:9999;display:flex;align-items:stretch;}
.sidebar-drawer{background:#161b22;border:1px solid #30363d;border-right:none;border-radius:10px 0 0 10px;width:0;overflow:hidden;opacity:0;transition:width 0.25s ease,opacity 0.2s ease;display:flex;flex-direction:column;}
.sidebar-drawer.open{width:220px;opacity:1;}
.sidebar-inner{padding:16px;white-space:nowrap;min-width:220px;}
.sidebar-title{font-size:10px;color:#484f58;text-transform:uppercase;letter-spacing:1.5px;margin-bottom:10px;font-family:'DM Mono',monospace;}
.sidebar-link{display:block;font-size:12px;color:#58a6ff;text-decoration:none;padding:7px 10px;border-radius:6px;margin-bottom:4px;line-height:1.4;transition:background 0.15s;}
.sidebar-link:hover{background:#1f4068;}
.sidebar-divider{border:none;border-top:1px solid #21262d;margin:8px 0;}
.sidebar-tab{background:#1f2937;color:#58a6ff;border:1px solid #30363d;border-right:none;border-radius:10px 0 0 10px;padding:12px 7px;cursor:pointer;display:flex;align-items:center;justify-content:center;width:28px;flex-shrink:0;transition:background 0.2s;}
.sidebar-tab:hover{background:#1f4068;}
.sidebar-tab span{writing-mode:vertical-rl;text-orientation:mixed;font-size:10px;font-weight:700;letter-spacing:2px;text-transform:uppercase;font-family:'DM Sans',sans-serif;user-select:none;}
</style></head>
<body>

<div class="sidebar-wrap" id="sidebarWrap">
    <div class="sidebar-drawer" id="sidebarDrawer">
        <div class="sidebar-inner">
            <div class="sidebar-title">Resources</div>
            <a class="sidebar-link" href="https://ico.org.uk/for-organisations/uk-gdpr-guidance-and-resources/" target="_blank" rel="noopener">Current UK GDPR Guidance</a>
            <a class="sidebar-link" href="https://ico.org.uk/about-the-ico/" target="_blank" rel="noopener">About the ICO</a>
            <hr class="sidebar-divider">
            <a class="sidebar-link" href="https://www.legislation.gov.uk/ukpga/2000/36/contents" target="_blank" rel="noopener">Freedom of Information Act 2000</a>
            <a class="sidebar-link" href="https://www.legislation.gov.uk/ukpga/2018/12/contents" target="_blank" rel="noopener">Data Protection Act 2018</a>
            <hr class="sidebar-divider">
            <a class="sidebar-link" href="https://www.rotherham.gov.uk/consultation-feedback/freedom-information-request-foi" target="_blank" rel="noopener">RMBC FOI Policy</a>
        </div>
    </div>
    <div class="sidebar-tab" onclick="toggleSidebar()"><span>Resources</span></div>
</div>
<script>
function toggleSidebar(){
    var d=document.getElementById('sidebarDrawer');
    if(d) d.classList.toggle('open');
}
</script>
<div class="topbar">
    <div>
        <div class="wordmark">CamScribe</div>
        <h1>🎥 RMBC CCTV — {% if pipeline == 'syp' %}SYP Witness Statement{% else %}RMBC Witness Statement{% endif %}</h1>
    </div>
    <div><a href="/">← Bookmarks</a><a href="/logout">Sign out</a></div>
</div>
<div class="container">

    <div class="incident">
        <h3>📋 Incident: {{ bm.name }}</h3>
        <p>
            Start &nbsp;&nbsp;&nbsp;&nbsp;: {{ bm.start_fmt }}<br>
            End &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;: {{ bm.end_fmt }}<br>
            Duration &nbsp;: {{ bm.duration_fmt }}<br>
            Bookmark created : {{ bm.created_fmt }}<br>
            Description : {{ bm.description or "None" }}
        </p>
    </div>

    <form id="stmtForm" method="POST">

        <!-- 1. YOUR DETAILS -->
        <div class="section">
            <h3>👤 Your Details</h3>
            <div class="row">
                <div><label>Full Name</label><input type="text" name="witness_name" value="{{ session.user_name }}" required></div>
                <div><label>Role / Job Title</label><input type="text" name="witness_role" value="{{ session.user_role }}" required></div>
            </div>
            <div class="row">
                <div><label>Organisation</label><input type="text" name="witness_org" value="Rotherham Metropolitan Borough Council" required></div>
                <div>
                    <label>Currently Based At</label>
                    <select name="witness_base" id="witnessBase">
                        {% for loc in locations %}
                        <option value="{{ loc }}">{{ loc }}</option>
                        {% endfor %}
                    </select>
                </div>
            </div>
            <div class="row">
                <div><label>Date of Statement</label><input type="text" name="statement_date" id="statementDate" value="{{ today }}" required></div>
            </div>
        </div>

        <!-- 2. INCIDENT REFERENCE -->
        <div class="section">
            <h3>📁 Incident Reference</h3>
            <div class="toggle-row">
                <input type="checkbox" id="no_ref_check" onchange="toggleRef(this)">
                <label for="no_ref_check">No reference available — skip this section</label>
            </div>
            <div id="ref_fields" style="margin-top:4px;">
                <div class="row">
                    <div><label>Incident / Crime Name</label><input type="text" name="crime_name" placeholder="e.g. Criminal Damage"></div>
                    {% if pipeline == 'syp' %}<div><label>Crime Reference Number <span>(if known)</span></label><input type="text" name="crime_number" placeholder="e.g. 22/12345/24"></div>{% endif %}
                </div>
                {% if pipeline != 'syp' %}
                <div class="row">
                    <div><label>Flare Reference</label><input type="text" name="flare_ref" placeholder="e.g. FLR-2026-001"></div>
                </div>
                {% endif %}
            </div>
        </div>

        <!-- 3. LOCATION & CAMERA -->
        <div class="section">
            <h3>📍 Location &amp; Camera</h3>
            <div class="row">
                <div><label>Incident Location / Site Name</label><input type="text" name="incident_location" placeholder="e.g. Johns Street, Eastwood, Rotherham" required></div>
                <div>
                    <label>Post / Mounting Position</label>
                    <input type="text" name="camera_location" id="cameraLocation" placeholder="Lamp Post — add number, e.g. Lamp Post 5" value="Lamp Post ">
                </div>
            </div>
            <div style="font-size:12px;color:#484f58;margin-top:6px;">If no post number: enter <em>a Lamp Post on [Road Name]</em></div>
        </div>

        <!-- 4. BOOKMARK CREATOR -->
        <div class="section">
            <h3>🔖 Bookmark Creator</h3>
            <p class="hint" style="margin-top:0;">Who created the bookmark in Nx Witness?</p>
            <div class="row">
                <div>
                    <label>Bookmark Created By</label>
                    <select name="bookmark_creator" id="bookmarkCreator">
                        <option value="{{ session.user_name }}">{{ session.user_name }} (me)</option>
                        <option value="__other__">Someone else — enter name below</option>
                    </select>
                </div>
                <div id="other_creator_field" style="display:none;">
                    <label>Creator's Full Name</label>
                    <input type="text" name="bookmark_creator_other" placeholder="Full name of person who created the bookmark">
                </div>
            </div>
        </div>

        <!-- 5. NVR CLOCK CHECK -->
        <div class="green-section">
            <h3>🕐 NVR Clock Check</h3>
            <div class="clock-box" style="margin-bottom:16px;">
                <div style="display:grid;grid-template-columns:1fr 1fr;gap:20px;align-items:center;">
                    <div>
                        <div class="clock-label">📅 UK Official Time (Browser)</div>
                        <div class="clock-date" id="liveDate">Loading...</div>
                        <div class="clock-time" id="liveTime">--:--:--</div>
                        <div class="clock-utc" id="tzLabel">Loading...</div>
                    </div>
                    <div>
                        <div class="clock-label">🖥️ NVR Server Time</div>
                        <div class="clock-date" id="serverDate">Loading...</div>
                        <div class="clock-time" style="font-size:36px;">
                            <span id="serverTimeHHMM">--:--:</span><span id="serverTimeSS" style="transition:color 1.5s ease;color:#58a6ff;">--</span>
                        </div>
                        <div class="clock-utc" id="serverDiff">Fetching...</div>
                    </div>
                </div>
                <button type="button" onclick="useTheseTimesForClockCheck()" style="margin-top:16px;background:#238636;color:white;border:none;border-radius:6px;padding:10px 24px;font-size:13px;font-weight:600;cursor:pointer;font-family:'DM Sans',sans-serif;">
                    ✅ Use These Times for Clock Check
                </button>
            </div>
            <div class="toggle-row">
                <input type="checkbox" id="clock_checked_box" name="clock_checked" value="yes">
                <label for="clock_checked_box">I have checked the NVR clock against the BT Speaking Clock</label>
            </div>
            <div id="clock_fields" style="display:none;margin-top:14px;">
                <div class="row">
                    <div><label>Date of Clock Check</label><input type="text" name="clock_check_date" id="clockDateField" value="{{ today }}"></div>
                    <div><label>BT Speaking Clock Time</label><input type="text" name="bt_clock_time" placeholder="e.g. 08:50:07"></div>
                </div>
                <div class="row">
                    <div><label>NVR System Displayed</label><input type="text" name="nvr_clock_time" placeholder="e.g. 08:45:07"></div>
                    <div><label>Difference</label><input type="text" name="clock_difference" placeholder="e.g. 8 seconds"></div>
                </div>
                <div>
                    <label>NVR Clock Is</label>
                    <select name="clock_fast_slow">
                        <option value="fast">Fast (NVR shows a later time than BT clock)</option>
                        <option value="slow">Slow (NVR shows an earlier time than BT clock)</option>
                        <option value="accurate">Accurate (within acceptable range)</option>
                    </select>
                </div>
            </div>
        </div>

        <!-- 6. EXPORT & CUSTODY -->
        <div class="section">
            <h3>🔗 Export &amp; Chain of Custody</h3>
            <div class="row">
                <div><label>Date Footage Exported</label><input type="text" name="export_date" id="exportDateField" value="{{ today }}"></div>
                <div><label>Time Footage Exported</label><input type="text" name="export_time" id="exportTimeField" placeholder="e.g. 09:30:00"></div>
            </div>
            <div class="row">
                <div><label>Exhibit Reference</label><input type="text" name="exhibit_ref" id="exhibitRef" value="{{ initials }}1" required></div>
                <div>
                    <label>Transferable Media Used</label>
                    <select name="media_type" id="mediaTypeSelect">
                        <option value="Raw Data">Raw Data</option>
                        <option value="DVD Disc">DVD Disc</option>
                        <option value="USB Device">USB Device</option>
                    </select>
                </div>
            </div>
        </div>

        <!-- 7. HANDOVER -->
        <div class="section">
            <h3>🤝 Handover Method</h3>

            {% if pipeline == 'syp' %}
            <!-- SYP: DEMS | Personal Handover | Secure Storage -->
            <div class="handover-opts" style="grid-template-columns:1fr 1fr 1fr;">
                <div class="h-opt" id="opt_dems" onclick="setHandover('dems')">
                    <input type="radio" name="handover_type" value="dems">
                    <div class="icon">💻</div>
                    <div class="title">DEMS Upload</div>
                    <div class="sub">SYP Digital Evidence System</div>
                </div>
                <div class="h-opt" id="opt_person" onclick="setHandover('person')">
                    <input type="radio" name="handover_type" value="person">
                    <div class="icon">👮</div>
                    <div class="title">Personal Handover</div>
                    <div class="sub">Handed directly to officer</div>
                </div>
                <div class="h-opt active" id="opt_storage" onclick="setHandover('storage')">
                    <input type="radio" name="handover_type" value="storage" checked>
                    <div class="icon">🔒</div>
                    <div class="title">Secure Storage Locker</div>
                    <div class="sub">RMBC storage — collection TBC</div>
                </div>
            </div>

            <div id="fields_dems" style="display:none;margin-top:14px;">
                <div style="font-size:12px;color:#58a6ff;font-weight:700;text-transform:uppercase;letter-spacing:0.5px;margin-bottom:12px;background:#0d1f2d;border:1px solid #1f4068;border-radius:8px;padding:12px;">
                    💻 Complete DEMS transfer details below — leave blank if not yet uploaded.
                </div>
                <div class="row">
                    <div><label>Date of Upload</label><input type="text" name="electronic_date" placeholder="e.g. 06/04/2026"></div>
                    <div><label>Time of Upload</label><input type="text" name="electronic_time" placeholder="e.g. 09:30"></div>
                </div>
                <div><label>Uploaded By</label><input type="text" name="electronic_by" value="{{ session.user_name }}"></div>
                <div><label>Recipient / Access Provided To</label><input type="text" name="electronic_recipient" placeholder="e.g. DS Smith, South Yorkshire Police"></div>
                <div><label>DEMS Reference / Job Number <span>(if known)</span></label><input type="text" name="electronic_ref" placeholder="Leave blank if not yet known"></div>
            </div>

            {% else %}
            <!-- RMBC: Secure Storage | Case File -->
            <div class="handover-opts" style="grid-template-columns:1fr 1fr;">
                <div class="h-opt active" id="opt_storage" onclick="setHandover('storage')">
                    <input type="radio" name="handover_type" value="storage" checked>
                    <div class="icon">🔒</div>
                    <div class="title">Secure Storage Locker</div>
                    <div class="sub">RMBC storage — collection TBC</div>
                </div>
                <div class="h-opt" id="opt_casefile" onclick="setHandover('casefile')">
                    <input type="radio" name="handover_type" value="casefile">
                    <div class="icon">🗂</div>
                    <div class="title">Case File</div>
                    <div class="sub">Stored on RMBC internal system</div>
                </div>
            </div>

            <div id="fields_casefile" style="display:none;margin-top:12px;">
                <p style="font-size:13px;color:#8b949e;padding:12px;background:#0d1117;border-radius:6px;border:1px solid #21262d;">
                    📁 The exhibit will be recorded as stored on the Flare case management system for internal council use.
                </p>
            </div>
            {% endif %}

            <!-- Shared: Personal Handover fields -->
            <div id="fields_person" style="display:none;margin-top:16px;">
                <div class="row">
                    <div><label>Officer Name</label><input type="text" name="officer_name" id="officerName" placeholder="Police Officer's Name"></div>
                    <div><label>Collar / Warrant Number</label><input type="text" name="officer_number" placeholder="e.g. 1234"></div>
                </div>
                <div class="row3">
                    <div><label>Handover Location</label><input type="text" name="handover_location" id="handoverLocation" value="Rawmarsh Police Station"></div>
                    <div><label>Handover Date</label><input type="text" name="handover_date" id="handoverDate" value="{{ today }}"></div>
                    <div><label>Handover Time</label><input type="text" name="handover_time" id="handoverTime" placeholder="e.g. 09:15"></div>
                </div>
            </div>

            <!-- Shared: Secure Storage fields -->
            <div id="fields_storage" style="margin-top:12px;">
                <p style="font-size:13px;color:#8b949e;padding:12px;background:#0d1117;border-radius:6px;border:1px solid #21262d;">
                    🔒 The exhibit will be recorded as held in a secure CCTV storage locker.
                </p>
            </div>
        </div>

        <!-- ── Cloud delivery options ──────────────────────────────────── -->
        {% if wasabi_found %}
        <div style="background:#0d1f2d;border:1px solid #1f4068;border-radius:10px;padding:20px 24px;margin-bottom:14px;">
            <div style="font-size:13px;font-weight:700;color:#58a6ff;margin-bottom:4px;">&#9729;&#65039; Statement Delivery</div>
            <div style="font-size:11px;color:#484f58;font-family:'DM Mono',monospace;margin-bottom:10px;">{% if '[C]' in (bm.name or '').upper() %}[C] instant cloud bookmark — footage confirmed at page load{% elif '[D]' in (bm.name or '').upper() %}[D] deferred bookmark — footage now confirmed in cloud{% else %}footage confirmed in cloud at page load{% endif %}</div>
            <div style="display:flex;align-items:center;gap:16px;margin-bottom:14px;flex-wrap:wrap;">
                <span style="font-size:12px;color:#3fb950;">&#10003; Package confirmed</span>
                {% if wasabi_arrived %}<span style="font-size:12px;color:#8b949e;font-family:'DM Mono',monospace;">{{ wasabi_arrived }}</span>{% endif %}
            </div>
            <div style="font-size:12px;color:#484f58;margin-bottom:14px;">Choose how to deliver this statement:</div>
            <div style="display:flex;flex-direction:column;gap:10px;">
                <label style="display:flex;align-items:center;gap:10px;cursor:pointer;margin:0;text-transform:none;letter-spacing:0;font-size:14px;color:#e6edf3;font-weight:400;">
                    <input type="radio" name="delivery" value="download" style="width:auto;"> &#11015;&#65039; Download only
                </label>
                <label style="display:flex;align-items:center;gap:10px;cursor:pointer;margin:0;text-transform:none;letter-spacing:0;font-size:14px;color:#e6edf3;font-weight:400;">
                    <input type="radio" name="delivery" value="cloud" style="width:auto;"> &#9729;&#65039; Save into cloud package only
                </label>
                <label style="display:flex;align-items:center;gap:10px;cursor:pointer;margin:0;text-transform:none;letter-spacing:0;font-size:14px;color:#e6edf3;font-weight:400;">
                    <input type="radio" name="delivery" value="both" checked style="width:auto;"> &#11015;&#65039;&#9729;&#65039; Download + save into cloud package
                </label>
            </div>
            <div style="font-size:11px;color:#484f58;margin-top:12px;line-height:1.5;">Saving into the cloud package retains this bookmark's footage beyond the normal overwrite cycle under the Wasabi retention workflow.</div>
            <input type="hidden" name="wasabi_prefix" value="{{ wasabi_prefix }}">
        </div>
        {% else %}
        <div style="background:#161b22;border:1px solid #30363d;border-radius:8px;padding:12px 16px;margin-bottom:14px;">
            {% if '[D]' in (bm.name or '').upper() %}
            <div style="font-size:12px;color:#8b949e;">&#8987; <strong style="color:#d29922;">[D] Deferred upload</strong> — footage not yet in cloud. Check again after 06:00. <strong style="color:#e6edf3;">Download only for now.</strong></div>
            {% elif '[C]' in (bm.name or '').upper() %}
            <div style="font-size:12px;color:#8b949e;">&#9888;&#65039; <strong style="color:#f85149;">[C] Cloud bookmark</strong> — footage not found in cloud. <strong style="color:#e6edf3;">Download only.</strong> Contact CCTV team if upload is expected.</div>
            {% else %}
            <div style="font-size:12px;color:#8b949e;">&#9729;&#65039; No footage found in cloud for this bookmark — <strong style="color:#e6edf3;">Download only</strong>.</div>
            {% endif %}
            <input type="hidden" name="delivery" value="download">
            <input type="hidden" name="wasabi_prefix" value="">
        </div>
        {% endif %}
        <!-- ───────────────────────────────────────────────────────────────── -->

        <!-- Officer Responsibility Agreement — after all delivery choices -->
        <div style="background:#1a1200;border:1px solid #d29922;border-radius:10px;padding:20px 24px;margin-bottom:14px;">
            <div style="font-size:13px;font-weight:700;color:#d29922;margin-bottom:10px;">Officer Responsibility</div>
            <p style="font-size:13px;color:#c9a227;line-height:1.7;margin-bottom:14px;">
                Once this footage is downloaded, it is the responsibility of the receiving officer to keep it secure and auditable in accordance with the security and accountability principles of UK GDPR (Article 5).
            </p>
            <label style="display:flex;align-items:flex-start;gap:10px;cursor:pointer;margin:0;text-transform:none;letter-spacing:0;font-size:13px;color:#d29922;font-weight:400;">
                <input type="checkbox" id="responsibilityCheck" onchange="toggleSubmit(this)" style="width:auto;margin-top:2px;flex-shrink:0;">
                I acknowledge and accept responsibility for the secure handling of this footage.
            </label>
        </div>

        <button type="submit" class="submit-btn" id="submitBtn" disabled style="opacity:0.5;cursor:not-allowed;">&#9889; Generate Witness Statement</button>

        <div id="loadingBox">
            <div style="font-size:36px;margin-bottom:12px;">⏳</div>
            <div style="font-size:16px;font-weight:700;color:#e6edf3;margin-bottom:8px;">Generating Statement...</div>
            <div style="font-size:13px;color:#8b949e;">Please wait — this may take up to 60 seconds.</div>
        </div>

    </form>
</div>

<div id="successBox" style="display:none;max-width:820px;margin:0 auto;padding:0 20px 80px;">
    <div style="background:#0d1117;border:1px solid #238636;border-radius:10px;overflow:hidden;">

        <!-- Header -->
        <div style="background:#0d2b0d;padding:24px 30px;text-align:center;border-bottom:1px solid #21262d;">
            <div style="font-size:36px;margin-bottom:8px;">&#10003;</div>
            <div style="font-size:18px;font-weight:700;color:#3fb950;margin-bottom:4px;">Statement Generated</div>
            <div id="successSubtitle" style="font-size:13px;color:#8b949e;">Your Word document is downloading now.</div>
        </div>

        <!-- Cloud timeline (shown only when uploaded) -->
        <div id="cloudTimeline" style="display:none;padding:14px 24px;border-bottom:1px solid #21262d;background:#0a1628;">
            <div style="font-size:10px;color:#484f58;text-transform:uppercase;letter-spacing:1.2px;margin-bottom:8px;font-family:'DM Mono',monospace;">Cloud Package Timeline</div>
            <div style="display:flex;flex-direction:column;gap:7px;">
                <div style="display:flex;gap:12px;align-items:baseline;">
                    <span style="font-size:11px;color:#484f58;font-family:'DM Mono',monospace;min-width:150px;flex-shrink:0;">&#128230; Package confirmed</span>
                    <span id="tlArrived" style="font-size:13px;color:#8b949e;font-family:'DM Mono',monospace;"></span>
                </div>
                <div style="display:flex;gap:12px;align-items:baseline;">
                    <span style="font-size:11px;color:#484f58;font-family:'DM Mono',monospace;min-width:150px;flex-shrink:0;">&#128196; Statement saved</span>
                    <span id="tlSaved" style="font-size:13px;color:#3fb950;font-family:'DM Mono',monospace;"></span>
                </div>
            </div>
        </div>
        <div id="cloudError" style="display:none;font-size:13px;color:#f85149;padding:12px 24px;border-bottom:1px solid #21262d;"></div>

        <!-- Email section -->
        <div style="padding:20px 24px;border-bottom:1px solid #21262d;">
            <div style="font-size:11px;color:#8b949e;text-transform:uppercase;letter-spacing:1px;margin-bottom:12px;">Send by Email</div>
            <div style="display:grid;grid-template-columns:1fr auto 1fr;gap:8px;align-items:end;">
                <div>
                    <div style="font-size:11px;color:#8b949e;margin-bottom:5px;text-transform:uppercase;letter-spacing:0.4px;">Recipient Name</div>
                    <input type="text" id="emailRecipName" placeholder="e.g. david.brown" style="width:100%;padding:9px 12px;background:#0d1117;border:1px solid #30363d;border-radius:6px;font-size:14px;color:#e6edf3;font-family:'DM Sans',sans-serif;">
                </div>
                <div style="padding-bottom:2px;font-size:18px;color:#484f58;text-align:center;">@</div>
                <div>
                    <div style="font-size:11px;color:#8b949e;margin-bottom:5px;text-transform:uppercase;letter-spacing:0.4px;">Domain</div>
                    <select id="emailDomain" style="width:100%;padding:9px 12px;background:#0d1117;border:1px solid #30363d;border-radius:6px;font-size:14px;color:#e6edf3;font-family:'DM Sans',sans-serif;">
                        <option value="rotherham.gov.uk">rotherham.gov.uk</option>
                        <option value="southyorkshire.police.uk">southyorkshire.police.uk</option>
                    </select>
                </div>
            </div>
            <div id="emailStatus" style="font-size:13px;margin-top:8px;min-height:18px;"></div>
            <button onclick="sendEmail()" id="sendEmailBtn" style="width:100%;margin-top:10px;padding:11px;background:#1f4068;color:#58a6ff;border:1px solid #1f4068;border-radius:6px;font-size:14px;font-weight:600;cursor:pointer;font-family:'DM Sans',sans-serif;">
                Send Statement by Email
            </button>
        </div>

        <!-- Footer: New Statement -->
        <div style="padding:16px 24px;text-align:center;">
            <a href="/" style="background:#21262d;color:#8b949e;padding:10px 20px;border-radius:6px;font-size:13px;font-weight:600;text-decoration:none;display:inline-block;">&#8592; New Statement</a>
        </div>

    </div>
</div>

<script>
var _wasabiArrived = "{{ wasabi_arrived }}";
function getUKTime() {
    const now = new Date();
    function lastSunday(year, month) {
        const d = new Date(Date.UTC(year, month, 31));
        d.setUTCDate(31 - d.getUTCDay());
        return d;
    }
    const year = now.getUTCFullYear();
    const isBST = now >= lastSunday(year, 2) && now < lastSunday(year, 9);
    const uk = new Date(now.getTime() + (isBST ? 1 : 0) * 3600000);
    const days   = ['Sunday','Monday','Tuesday','Wednesday','Thursday','Friday','Saturday'];
    const months = ['January','February','March','April','May','June','July','August','September','October','November','December'];
    var hh = String(uk.getUTCHours()).padStart(2,'0');
    var mm = String(uk.getUTCMinutes()).padStart(2,'0');
    var ss = String(uk.getUTCSeconds()).padStart(2,'0');
    document.getElementById('liveDate').textContent = days[uk.getUTCDay()] + ' ' + uk.getUTCDate() + ' ' + months[uk.getUTCMonth()] + ' ' + uk.getUTCFullYear();
    document.getElementById('liveTime').textContent = hh + ':' + mm + ':' + ss;
    document.getElementById('tzLabel').textContent  = (isBST ? 'British Summer Time (BST)' : 'Greenwich Mean Time (GMT)') + ' · Updates every second';
    var et = document.getElementById('exportTimeField');
    if (et && !et.dataset.touched && !et.value) et.placeholder = hh + ':' + mm + ':' + ss + ' (now)';
    var ht = document.getElementById('handoverTime');
    if (ht && !ht.dataset.touched && !ht.value) ht.placeholder = hh + ':' + mm + ' (now)';
}
getUKTime();
setInterval(getUKTime, 1000);

function fetchServerTime() {
    fetch('/server-time')
        .then(function(r) { return r.json(); })
        .then(function(data) {
            document.getElementById('serverDate').textContent = data.date;
            var parts   = data.time.split(':');
            var hhmmStr = parts[0] + ':' + parts[1] + ':';
            var secStr  = parts[2];
            document.getElementById('serverTimeHHMM').textContent = hhmmStr;
            var secEl = document.getElementById('serverTimeSS');
            secEl.textContent = secStr;
            secEl.style.transition = 'none';
            secEl.style.color = '#3fb950';
            secEl.style.textShadow = '0 0 12px #3fb950';
            setTimeout(function() {
                secEl.style.transition = 'color 1.5s ease, text-shadow 1.5s ease';
                secEl.style.color = '#58a6ff';
                secEl.style.textShadow = 'none';
            }, 200);
            var btText = document.getElementById('liveTime').textContent;
            if (btText && btText !== '--:--:--') {
                var b = btText.split(':').map(Number);
                var s = data.time.split(':').map(Number);
                var diffSecs = (s[0]*3600 + s[1]*60 + s[2]) - (b[0]*3600 + b[1]*60 + b[2]);
                var absDiff  = Math.abs(diffSecs);
                var dMins = Math.floor(absDiff / 60);
                var dSecs = absDiff % 60;
                var diffStr;
                if (absDiff < 10)   diffStr = 'Within 10 seconds — accurate';
                else if (dMins > 0) diffStr = 'Server is ' + (diffSecs > 0 ? 'fast' : 'slow') + ' by ' + dMins + 'm ' + dSecs + 's';
                else                diffStr = 'Server is ' + (diffSecs > 0 ? 'fast' : 'slow') + ' by ' + dSecs + ' seconds';
                document.getElementById('serverDiff').textContent = diffStr;
            }
        })
        .catch(function() {
            document.getElementById('serverDate').textContent = 'Unavailable';
            document.getElementById('serverTimeHHMM').textContent = '--:--:';
            document.getElementById('serverTimeSS').textContent  = '--';
            document.getElementById('serverDiff').textContent    = 'Could not reach server';
        });
}
fetchServerTime();
setInterval(fetchServerTime, 5000);

function useTheseTimesForClockCheck() {
    var btTime  = document.getElementById('liveTime').textContent;
    var hhmm    = document.getElementById('serverTimeHHMM').textContent;
    var secPart = document.getElementById('serverTimeSS').textContent;
    var srvTime = hhmm + secPart;
    if (btTime === '--:--:--' || srvTime === '--:--:--') { alert('Times not yet loaded — wait a moment and try again.'); return; }
    var cb = document.getElementById('clock_checked_box');
    if (cb) { cb.checked = true; toggleClock(cb); }
    var cd = document.querySelector('[name="clock_check_date"]');
    var bf = document.querySelector('[name="bt_clock_time"]');
    var nf = document.querySelector('[name="nvr_clock_time"]');
    var df = document.querySelector('[name="clock_difference"]');
    var fs = document.querySelector('[name="clock_fast_slow"]');
    var sd = document.getElementById('statementDate');
    if (cd) cd.value = sd ? sd.value : document.getElementById('serverDate').textContent;
    if (bf) bf.value = btTime;
    if (nf) nf.value = srvTime;
    var b = btTime.split(':').map(Number);
    var s = srvTime.split(':').map(Number);
    var diffSecs = (s[0]*3600 + s[1]*60 + s[2]) - (b[0]*3600 + b[1]*60 + b[2]);
    var absDiff  = Math.abs(diffSecs);
    var dMins = Math.floor(absDiff / 60);
    var dSecs = absDiff % 60;
    if (df) {
        if (dMins > 0) df.value = dMins + ' minute' + (dMins !== 1 ? 's' : '') + ' and ' + dSecs + ' second' + (dSecs !== 1 ? 's' : '');
        else           df.value = absDiff + ' second' + (absDiff !== 1 ? 's' : '');
    }
    if (fs) {
        if (absDiff === 0)      fs.value = 'accurate';
        else if (diffSecs > 0)  fs.value = 'fast';
        else                    fs.value = 'slow';
    }
    var gs = document.querySelector('.green-section');
    if (gs) gs.scrollIntoView({behavior:'smooth', block:'center'});
}

function setHandover(type) {
    ['person','dems','storage','casefile'].forEach(function(t) {
        var opt = document.getElementById('opt_' + t);
        if (opt) opt.classList.remove('active');
        var f = document.getElementById('fields_' + t);
        if (f) f.style.display = 'none';
    });
    var activeOpt = document.getElementById('opt_' + type);
    if (activeOpt) activeOpt.classList.add('active');
    var radio = document.querySelector('input[name="handover_type"][value="' + type + '"]');
    if (radio) radio.checked = true;
    var show = document.getElementById('fields_' + type);
    if (show) show.style.display = 'block';
    if (type === 'person') { setTimeout(function() { var n = document.getElementById('officerName'); if (n) n.focus(); }, 100); }
    // DEMS is a digital route — Raw Data is the correct media type
    var mt = document.getElementById('mediaTypeSelect');
    if (mt && type === 'dems') mt.value = 'Raw Data';
}
setHandover('storage');

function toggleClock(cb) {
    var f = document.getElementById('clock_fields');
    if (f) f.style.display = cb.checked ? 'block' : 'none';
}
var clockCb = document.getElementById('clock_checked_box');
if (clockCb) clockCb.addEventListener('change', function() { toggleClock(this); });

function toggleRef(cb) {
    var f = document.getElementById('ref_fields');
    if (f) f.style.display = cb.checked ? 'none' : 'block';
}

// Cloud delivery is a digital route — auto-select Raw Data as media type
document.querySelectorAll('input[name="delivery"]').forEach(function(r) {
    r.addEventListener('change', function() {
        var mt = document.getElementById('mediaTypeSelect');
        if (mt && (this.value === 'cloud' || this.value === 'both')) mt.value = 'Raw Data';
    });
});

var bcEl = document.getElementById('bookmarkCreator');
if (bcEl) bcEl.addEventListener('change', function() {
    var f = document.getElementById('other_creator_field');
    if (f) f.style.display = this.value === '__other__' ? 'block' : 'none';
});

var allInputs = document.querySelectorAll('input');
for (var i = 0; i < allInputs.length; i++) {
    allInputs[i].addEventListener('input', function() { this.dataset.touched = '1'; });
}

function toggleSubmit(cb) {
    var btn = document.getElementById('submitBtn');
    if (!btn) return;
    if (cb.checked) {
        btn.disabled = false;
        btn.style.opacity = '1';
        btn.style.cursor = 'pointer';
    } else {
        btn.disabled = true;
        btn.style.opacity = '0.5';
        btn.style.cursor = 'not-allowed';
    }
}

var stmtForm = document.getElementById('stmtForm');
var _docToken = null;
if (stmtForm) stmtForm.addEventListener('submit', function(e) {
    e.preventDefault();
    var creatorSel = document.getElementById('bookmarkCreator');
    if (creatorSel && creatorSel.value === '__other__') {
        var otherInput = document.querySelector('input[name="bookmark_creator_other"]');
        creatorSel.value = (otherInput && otherInput.value.trim()) ? otherInput.value.trim() : 'Unknown';
    }
    var btn = document.getElementById('submitBtn');
    btn.disabled = true; btn.textContent = '⏳ Generating...';
    document.getElementById('loadingBox').style.display = 'block';
    fetch(window.location.href, {method:'POST', body: new FormData(this)})
        .then(function(r) {
            if (!r.ok) throw new Error('Server error ' + r.status);
            _docToken = r.headers.get('X-Doc-Token');
            window._delivery       = r.headers.get('X-Delivery') || 'download';
            window._wasabiUploaded = r.headers.get('X-Wasabi-Uploaded') === 'true';
            window._wasabiTs       = r.headers.get('X-Wasabi-Timestamp') || '';
            window._wasabiErr      = r.headers.get('X-Wasabi-Error') || '';
            return r.blob();
        })
        .then(function(blob) {
            var delivery = window._delivery || 'download';
            if (delivery === 'download' || delivery === 'both') {
                var url = URL.createObjectURL(blob);
                var a = document.createElement('a'); a.href = url;
                a.download = 'witness_statement.docx';
                document.body.appendChild(a); a.click();
                document.body.removeChild(a); URL.revokeObjectURL(url);
            }
            document.getElementById('loadingBox').style.display = 'none';
            document.querySelector('.container').style.display  = 'none';
            document.getElementById('successBox').style.display = 'block';
            // Subtitle
            var sub = document.getElementById('successSubtitle');
            if (delivery === 'cloud') sub.textContent = 'Statement saved into cloud package — no local download.';
            else if (delivery === 'both') sub.textContent = 'Statement downloading and saved into cloud package.';
            else sub.textContent = 'Your Word document is downloading now.';
            // Cloud package timeline
            if (window._wasabiUploaded) {
                var tl = document.getElementById('cloudTimeline');
                if (tl) {
                    document.getElementById('tlArrived').textContent = _wasabiArrived || '—';
                    document.getElementById('tlSaved').textContent   = window._wasabiTs || '—';
                    tl.style.display = 'block';
                }
            } else if (window._wasabiErr) {
                var ce = document.getElementById('cloudError');
                if (ce) {
                    ce.textContent = '⚠ Statement upload to cloud failed: ' + window._wasabiErr;
                    ce.style.display = 'block';
                }
            }
        })
        .catch(function(err) {
            document.getElementById('loadingBox').innerHTML =
                '<div style="color:#f85149;font-size:15px;margin-bottom:12px;">❌ Error: ' + err.message + '</div>' +
                '<button onclick="location.reload()" style="background:#21262d;color:#e6edf3;padding:10px 20px;border-radius:6px;border:none;cursor:pointer;">Try Again</button>';
        });
});

function sendEmail() {
    var name   = document.getElementById('emailRecipName').value.trim();
    var domain = document.getElementById('emailDomain').value;
    var status = document.getElementById('emailStatus');
    var btn    = document.getElementById('sendEmailBtn');
    if (!name) { status.style.color='#f85149'; status.textContent='Please enter a recipient name.'; return; }
    if (!_docToken) { status.style.color='#f85149'; status.textContent='No document token — please regenerate the statement.'; return; }
    btn.disabled = true; btn.textContent = '⏳ Sending...';
    status.style.color = '#8b949e'; status.textContent = 'Sending to ' + name + '@' + domain + '...';
    fetch('/send-email', {
        method: 'POST',
        headers: {'Content-Type': 'application/json'},
        body: JSON.stringify({token: _docToken, recipient_name: name, recipient_domain: domain})
    })
    .then(function(r) { return r.json(); })
    .then(function(data) {
        if (data.ok) {
            status.style.color = '#3fb950';
            status.textContent = '✅ Email sent successfully to ' + data.recipient;
            btn.textContent = '📧 Send Another Copy';
            btn.disabled = false;
        } else {
            status.style.color = '#f85149';
            status.textContent = '❌ ' + data.error;
            btn.textContent = '📧 Send Statement by Email';
            btn.disabled = false;
        }
    })
    .catch(function(err) {
        status.style.color = '#f85149';
        status.textContent = '❌ Network error: ' + err.message;
        btn.textContent = '📧 Send Statement by Email';
        btn.disabled = false;
    });
}
</script>
</body></html>"""

FOI_FORM_HTML = """<!DOCTYPE html>
<html><head><title>RMBC CCTV — FOI Disclosure Record</title>
<meta name="viewport" content="width=device-width, initial-scale=1">
<link href="https://fonts.googleapis.com/css2?family=DM+Mono:wght@400;500&family=DM+Sans:wght@400;500;700&display=swap" rel="stylesheet">
<style>
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:'DM Sans',sans-serif;background:#0d1117;color:#e6edf3;}
.topbar{background:#161b22;border-bottom:1px solid #30363d;padding:14px 30px;display:flex;justify-content:space-between;align-items:center;}
.topbar h1{font-size:16px;font-weight:700;}
.topbar a{color:#58a6ff;text-decoration:none;font-size:13px;margin-left:16px;}
.wordmark{font-size:10px;font-weight:700;letter-spacing:3px;color:#58a6ff;font-family:'DM Mono',monospace;text-transform:uppercase;line-height:1;margin-bottom:3px;}
.container{max-width:820px;margin:30px auto;padding:0 20px 80px;}
.incident{background:#161b22;border:1px solid #30363d;border-radius:10px;padding:18px 22px;margin-bottom:20px;border-left:3px solid #8a5cf6;}
.incident h3{font-size:13px;font-weight:700;color:#8a5cf6;margin-bottom:8px;text-transform:uppercase;letter-spacing:0.5px;}
.incident p{font-size:13px;color:#8b949e;line-height:1.8;font-family:'DM Mono',monospace;}
.section{background:#161b22;border:1px solid #30363d;border-radius:10px;padding:24px;margin-bottom:14px;}
.section h3{font-size:12px;font-weight:700;color:#8b949e;text-transform:uppercase;letter-spacing:1px;margin-bottom:18px;padding-bottom:10px;border-bottom:1px solid #21262d;}
.section h3 span{font-size:10px;background:#1f2937;color:#484f58;padding:2px 8px;border-radius:4px;margin-left:8px;font-weight:400;text-transform:none;letter-spacing:0;}
.purple-section{background:#120a1f;border:1px solid #6e40c9;border-radius:10px;padding:24px;margin-bottom:14px;}
.purple-section h3{font-size:12px;font-weight:700;color:#8a5cf6;text-transform:uppercase;letter-spacing:1px;margin-bottom:18px;padding-bottom:10px;border-bottom:1px solid #2a1a4a;}
.green-section{background:#0a1f0a;border:1px solid #238636;border-radius:10px;padding:24px;margin-bottom:14px;}
.green-section h3{font-size:12px;font-weight:700;color:#3fb950;text-transform:uppercase;letter-spacing:1px;margin-bottom:18px;padding-bottom:10px;border-bottom:1px solid #1a3a1a;}
.warning-box{background:#2d1f00;border:1px solid #d29922;border-radius:8px;padding:12px 16px;margin-top:10px;font-size:12px;color:#d29922;line-height:1.6;}
.info-box{background:#0d1f2d;border:1px solid #1f4068;border-radius:8px;padding:12px 16px;margin-top:10px;font-size:12px;color:#58a6ff;line-height:1.6;}
label{display:block;font-size:12px;color:#8b949e;font-weight:500;margin-bottom:6px;margin-top:14px;text-transform:uppercase;letter-spacing:0.4px;}
label span{font-weight:400;text-transform:none;color:#484f58;margin-left:4px;}
input,select,textarea{width:100%;padding:10px 14px;background:#0d1117;border:1px solid #30363d;border-radius:6px;font-size:14px;color:#e6edf3;font-family:'DM Sans',sans-serif;}
input:focus,select:focus,textarea:focus{outline:none;border-color:#8a5cf6;}
select option{background:#161b22;}
textarea{resize:vertical;min-height:70px;}
.row{display:grid;grid-template-columns:1fr 1fr;gap:14px;}
.row3{display:grid;grid-template-columns:1fr 1fr 1fr;gap:14px;}
.toggle-row{display:flex;align-items:center;gap:12px;margin-top:10px;}
.toggle-row input[type=checkbox]{width:auto;margin:0;}
.toggle-row label{margin:0;text-transform:none;font-size:13px;color:#8b949e;letter-spacing:0;}
.submit-btn{width:100%;padding:16px;background:#6e40c9;color:white;border:none;border-radius:8px;font-size:16px;font-weight:700;cursor:pointer;margin-top:10px;font-family:'DM Sans',sans-serif;}
.submit-btn:hover{background:#8a5cf6;}
.submit-btn:disabled{background:#21262d;color:#484f58;cursor:not-allowed;}
#loadingBox{display:none;background:#161b22;border:1px solid #30363d;border-radius:10px;padding:30px;text-align:center;margin-top:16px;}
#successBox{display:none;}
.legal-badge{display:inline-block;background:#1a0d2b;color:#8a5cf6;border:1px solid #6e40c9;padding:6px 14px;border-radius:6px;font-size:12px;font-family:'DM Mono',monospace;margin-top:8px;}
@media(max-width:640px){
  .row,.row3{grid-template-columns:1fr !important;}
  .container{padding:0 12px 60px;}
  .topbar h1{font-size:13px;}
  .sidebar-wrap{top:auto;bottom:16px;transform:none;right:16px;flex-direction:column-reverse;align-items:flex-end;}
  .sidebar-drawer{border-radius:10px;border-right:1px solid #6e40c9;width:0;}
  .sidebar-drawer.open{width:220px;margin-bottom:8px;}
  .sidebar-tab{border-radius:50px;border-right:1px solid #6e40c9;width:auto;padding:8px 14px;}
  .sidebar-tab span{writing-mode:horizontal-tb;letter-spacing:1px;font-size:11px;}
}
.sidebar-wrap{position:fixed;right:0;top:50vh;transform:translateY(-50%);z-index:9999;display:flex;align-items:stretch;}
.sidebar-drawer{background:#161b22;border:1px solid #30363d;border-right:none;border-radius:10px 0 0 10px;width:0;overflow:hidden;opacity:0;transition:width 0.25s ease,opacity 0.2s ease;display:flex;flex-direction:column;}
.sidebar-drawer.open{width:220px;opacity:1;}
.sidebar-inner{padding:16px;white-space:nowrap;min-width:220px;}
.sidebar-title{font-size:10px;color:#484f58;text-transform:uppercase;letter-spacing:1.5px;margin-bottom:10px;font-family:'DM Mono',monospace;}
.sidebar-link{display:block;font-size:12px;color:#8a5cf6;text-decoration:none;padding:7px 10px;border-radius:6px;margin-bottom:4px;line-height:1.4;transition:background 0.15s;}
.sidebar-link:hover{background:#2a1a4a;}
.sidebar-divider{border:none;border-top:1px solid #21262d;margin:8px 0;}
.sidebar-tab{background:#2a1a4a;color:#8a5cf6;border:1px solid #6e40c9;border-right:none;border-radius:10px 0 0 10px;padding:12px 7px;cursor:pointer;display:flex;align-items:center;justify-content:center;width:28px;flex-shrink:0;transition:background 0.2s;}
.sidebar-tab:hover{background:#6e40c9;}
.sidebar-tab span{writing-mode:vertical-rl;text-orientation:mixed;font-size:10px;font-weight:700;letter-spacing:2px;text-transform:uppercase;font-family:'DM Sans',sans-serif;user-select:none;}
@media(max-width:768px){
  .sidebar-wrap{top:auto;bottom:16px;transform:none;right:16px;flex-direction:column-reverse;align-items:flex-end;}
  .sidebar-drawer{border-radius:10px;border-right:1px solid #6e40c9;width:0;}
  .sidebar-drawer.open{width:220px;margin-bottom:8px;}
  .sidebar-tab{border-radius:50px;border-right:1px solid #6e40c9;width:auto;padding:8px 14px;}
  .sidebar-tab span{writing-mode:horizontal-tb;letter-spacing:1px;font-size:11px;}
}
</style></head>
<body>

<div class="sidebar-wrap" id="sidebarWrap">
    <div class="sidebar-drawer" id="sidebarDrawer">
        <div class="sidebar-inner">
            <div class="sidebar-title">Resources</div>
            <a class="sidebar-link" href="https://ico.org.uk/for-organisations/uk-gdpr-guidance-and-resources/" target="_blank" rel="noopener">Current UK GDPR Guidance</a>
            <a class="sidebar-link" href="https://ico.org.uk/about-the-ico/" target="_blank" rel="noopener">About the ICO</a>
            <hr class="sidebar-divider">
            <a class="sidebar-link" href="https://www.legislation.gov.uk/ukpga/2000/36/contents" target="_blank" rel="noopener">Freedom of Information Act 2000</a>
            <a class="sidebar-link" href="https://www.legislation.gov.uk/ukpga/2018/12/contents" target="_blank" rel="noopener">Data Protection Act 2018</a>
            <hr class="sidebar-divider">
            <a class="sidebar-link" href="https://www.rotherham.gov.uk/consultation-feedback/freedom-information-request-foi" target="_blank" rel="noopener">RMBC FOI Policy</a>
        </div>
    </div>
    <div class="sidebar-tab" onclick="toggleSidebar()"><span>Resources</span></div>
</div>
<script>
function toggleSidebar(){
    var d=document.getElementById('sidebarDrawer');
    if(d) d.classList.toggle('open');
}
</script>
<div class="topbar">
    <div>
        <div class="wordmark">CamScribe</div>
        <h1>📋 RMBC CCTV — Disclosure Record</h1>
    </div>
    <div><a href="/">← Bookmarks</a><a href="/logout">Sign out</a></div>
</div>
<div class="container">

    <div class="incident">
        <h3>📋 Footage: {{ bm.name }}</h3>
        <p>Period &nbsp;&nbsp;: {{ bm.start_fmt }} → {{ bm.end_fmt }}<br>Duration : {{ bm.duration_fmt }}<br>System &nbsp;&nbsp;: {{ site_ref }}</p>
    </div>

    <form id="foiForm" method="POST">

        <!-- S1+2: Reference & Request Details -->
        <div class="section">
            <h3>📁 Section 1–2 · Reference &amp; Request Details</h3>
            <div class="row">
                <div><label>Disclosure / FOI Reference <span>(if known)</span></label><input type="text" name="foi_ref" placeholder="e.g. FOI-2026-001"></div>
                <div>
                    <label>Request Type</label>
                    <select name="foi_request_type" id="foiRequestType" onchange="updateLegalBasis()">
                        <option value="Public">Public (FOIA 2000)</option>
                        <option value="Solicitor">Solicitor / Legal</option>
                        <option value="Insurance">Insurance Company</option>
                    </select>
                </div>
            </div>
            <div class="row">
                <div><label>Requestor Name</label><input type="text" name="foi_requester" placeholder="Full name" required></div>
                <div><label>Requesting Organisation</label><input type="text" name="foi_organisation" placeholder="e.g. South Yorkshire Police"></div>
            </div>
            <div class="row">
                <div><label>Date Request Received</label><input type="text" name="foi_date_received" value="{{ today }}"></div>
                <div><label>Incident Type</label><input type="text" name="foi_incident_type" placeholder="e.g. Fly Tipping, Criminal Damage"></div>
            </div>
            <div><label>Summary of Request <span>(brief description)</span></label>
                <textarea name="foi_summary" placeholder="e.g. Request for footage of fly tipping incident at the above location." required></textarea>
            </div>
            <div><label>Incident Location</label><input type="text" name="incident_location" placeholder="e.g. High Street, Rotherham" required></div>
        </div>

        <!-- S3: Legal Basis -->
        <div class="purple-section">
            <h3>Section 3 · Legal Basis for Disclosure</h3>
            <p style="font-size:12px;color:#8b949e;margin-bottom:16px;line-height:1.6;">Disclosure is made in line with the Freedom of Information Act 2000, the Data Protection Act 2018, and UK GDPR.</p>

            <label style="font-size:12px;color:#8b949e;font-weight:500;text-transform:uppercase;letter-spacing:0.4px;margin-bottom:10px;display:block;">Footage contains identifiable individuals or vehicles</label>
            <div style="display:flex;gap:24px;margin-bottom:12px;">
                <label style="display:flex;align-items:center;gap:8px;margin:0;text-transform:none;font-size:13px;color:#e6edf3;letter-spacing:0;cursor:pointer;">
                    <input type="radio" id="identifiableYes" name="foi_identifiable" value="yes" onchange="updateLegalBasis()" style="width:auto;margin:0;">
                    Yes
                </label>
                <label style="display:flex;align-items:center;gap:8px;margin:0;text-transform:none;font-size:13px;color:#e6edf3;letter-spacing:0;cursor:pointer;">
                    <input type="radio" id="identifiableNo" name="foi_identifiable" value="no" onchange="updateLegalBasis()" checked style="width:auto;margin:0;">
                    No
                </label>
            </div>

            <div class="warning-box" id="identifiableWarning" style="display:none;">
                <strong>⚠️ Advisory</strong>
                Identifiable people or vehicle number plates are present — this disclosure must be processed under <strong>UK GDPR / DPA 2018</strong>, not FOIA 2000. Ensure the appropriate lawful basis is in place before disclosure.
            </div>
            <div class="legal-badge" id="legalBadge" style="margin-top:10px;">FOIA 2000</div>
            <input type="hidden" name="foi_legal_display" id="legalDisplay" value="FOIA 2000">
        </div>

        <!-- S4-5: Incident & Footage (auto from bookmark) -->
        <div class="section">
            <h3>📍 Section 4–5 · Incident &amp; Footage Details <span>auto-populated from bookmark</span></h3>
            <div class="info-box">
                ℹ️ Footage period, duration and system reference are automatically populated from the selected bookmark.
            </div>
            <div class="row" style="margin-top:14px;">
                <div><label>Footage Start</label><input type="text" value="{{ bm.start_fmt }}" readonly style="color:#484f58;cursor:not-allowed;"></div>
                <div><label>Footage End</label><input type="text" value="{{ bm.end_fmt }}" readonly style="color:#484f58;cursor:not-allowed;"></div>
            </div>
            <div class="row">
                <div><label>Duration</label><input type="text" value="{{ bm.duration_fmt }}" readonly style="color:#484f58;cursor:not-allowed;"></div>
                <div><label>System Reference</label><input type="text" value="{{ site_ref }}" readonly style="color:#484f58;cursor:not-allowed;"></div>
            </div>
        </div>

        <!-- 6: Format, Delivery & Exhibit -->
        <div class="section">
            <h3>💿 Section 6 · Format, Delivery &amp; Exhibit</h3>
            <div class="row">
                <div>
                    <label>Export Format</label>
                    <select name="foi_export_format">
                        <option value="MP4">MP4</option>
                        <option value="AVI">AVI</option>
                        <option value="Native NVR Format">Native NVR Format</option>
                    </select>
                </div>
                <div>
                    <label>Delivery Method</label>
                    <select name="media_type">
                        <option value="Raw Data Files">Raw Data Files</option>
                        <option value="DVD Disc">DVD Disc</option>
                        <option value="USB Device">USB Device</option>
                    </select>
                </div>
            </div>
            <div class="row">
                <div>
                    <label>Encryption Applied</label>
                    <select name="foi_encryption">
                        <option value="No">No</option>
                        <option value="Yes">Yes</option>
                    </select>
                </div>
                <div>
                    <label>Viewing Software Provided</label>
                    <select name="foi_viewing_software">
                        <option value="No">No</option>
                        <option value="Yes — VLC Media Player">Yes — VLC Media Player</option>
                        <option value="Yes — Nx Witness Client">Yes — Nx Witness Client</option>
                    </select>
                </div>
            </div>
            <div style="margin-top:14px;">
                <label>Exhibit ID</label>
                <input type="text" name="exhibit_ref" value="{{ initials }}1" required>
            </div>
        </div>

        <!-- 7: Redaction -->
        <div class="purple-section">
            <h3>Section 7 · Redaction Guidance</h3>
            <p style="font-size:12px;color:#8b949e;margin-bottom:12px;line-height:1.6;">This is the engineer's assessment only. The FOI Team / Information Governance makes the final redaction decision.</p>
            <div>
                <label>Does footage require redacting?</label>
                <select name="redaction_onsite" id="redactionOnsite">
                    <option value="yes">Yes — footage likely contains content requiring redaction</option>
                    <option value="no">No — footage appears suitable for disclosure without redaction</option>
                </select>
            </div>
        </div>

        <!-- 8: Time Verification (optional) -->
        <div class="green-section">
            <h3>🕐 Section 8 · Time Synchronisation Verification <span>optional</span></h3>
            <div style="background:#0d1117;border:1px solid #2a1a4a;border-radius:8px;padding:16px 20px;margin-bottom:16px;text-align:center;">
                <div style="display:grid;grid-template-columns:1fr 1fr;gap:20px;align-items:center;">
                    <div>
                        <div style="font-size:11px;color:#484f58;text-transform:uppercase;letter-spacing:1px;margin-bottom:8px;font-family:'DM Mono',monospace;">UK Official Time (Browser)</div>
                        <div style="font-size:15px;font-weight:600;color:#8b949e;margin-bottom:4px;font-family:'DM Mono',monospace;" id="foiLiveDate">Loading...</div>
                        <div style="font-size:36px;font-weight:700;color:#8a5cf6;letter-spacing:3px;font-family:'DM Mono',monospace;line-height:1;" id="foiLiveTime">--:--:--</div>
                        <div style="font-size:11px;color:#484f58;margin-top:6px;font-family:'DM Mono',monospace;" id="foiTzLabel">Loading...</div>
                    </div>
                    <div>
                        <div style="font-size:11px;color:#484f58;text-transform:uppercase;letter-spacing:1px;margin-bottom:8px;font-family:'DM Mono',monospace;">NVR Server Time</div>
                        <div style="font-size:15px;font-weight:600;color:#8b949e;margin-bottom:4px;font-family:'DM Mono',monospace;" id="foiServerDate">Loading...</div>
                        <div style="font-size:36px;font-weight:700;letter-spacing:3px;font-family:'DM Mono',monospace;line-height:1;">
                            <span id="foiServerHHMM" style="color:#8a5cf6;">--:--:</span><span id="foiServerSS" style="color:#8a5cf6;">--</span>
                        </div>
                        <div style="font-size:11px;color:#484f58;margin-top:6px;font-family:'DM Mono',monospace;" id="foiServerDiff">Fetching...</div>
                    </div>
                </div>
                <button type="button" onclick="foiUseTheseTimes()" style="margin-top:14px;background:#6e40c9;color:white;border:none;border-radius:6px;padding:9px 22px;font-size:13px;font-weight:600;cursor:pointer;font-family:'DM Sans',sans-serif;">
                    Use These Times for Clock Check
                </button>
            </div>
            <div class="toggle-row">
                <input type="checkbox" id="timeVerifyCheck" name="time_verified" value="yes" onchange="toggleTimeVerify(this)">
                <label for="timeVerifyCheck">Time verification was performed against BT Speaking Clock</label>
            </div>
            <div id="timeVerifyFields" style="display:none;margin-top:14px;">
                <div class="row">
                    <div><label>BT Speaking Clock Time</label><input type="text" name="foi_verify_time" placeholder="e.g. 09:30:00"></div>
                    <div><label>NVR System Time at Verification</label><input type="text" name="foi_system_time" placeholder="e.g. 09:29:45"></div>
                </div>
                <div><label>Offset / Difference</label><input type="text" name="foi_time_offset" placeholder="e.g. 15 seconds slow"></div>
                <div class="info-box" style="margin-top:10px;">ℹ️ Date/Time verification will be auto-populated within the statement if synced.</div>
            </div>
        </div>

        <!-- 9: Authorisation -->
        <div class="section">
            <h3>👤 Section 9 · Authorisation — Completed By</h3>
            <div class="row">
                <div><label>Full Name</label><input type="text" name="witness_name" value="{{ session.user_name }}" required></div>
                <div><label>Role</label><input type="text" name="witness_role" value="{{ session.user_role }}" required></div>
            </div>
            <div class="row">
                <div><label>Date</label><input type="text" name="statement_date" value="{{ today }}" required></div>
                <div><label>Contact</label><input type="text" name="witness_contact" placeholder="Email or extension"></div>
            </div>
        </div>

        <!-- ── Cloud delivery options (FOI) ───────────────────────────────── -->
        {% if wasabi_found %}
        <div style="background:#0d1f2d;border:1px solid #1f4068;border-radius:10px;padding:20px 24px;margin-bottom:14px;">
            <div style="font-size:13px;font-weight:700;color:#58a6ff;margin-bottom:4px;">&#9729;&#65039; Disclosure Record Delivery</div>
            <div style="font-size:11px;color:#484f58;font-family:'DM Mono',monospace;margin-bottom:10px;">{% if '[C]' in (bm.name or '').upper() %}[C] instant cloud bookmark — footage confirmed at page load{% elif '[D]' in (bm.name or '').upper() %}[D] deferred bookmark — footage now confirmed in cloud{% else %}footage confirmed in cloud at page load{% endif %}</div>
            <div style="display:flex;align-items:center;gap:16px;margin-bottom:14px;flex-wrap:wrap;">
                <span style="font-size:12px;color:#3fb950;">&#10003; Package confirmed</span>
                {% if wasabi_arrived %}<span style="font-size:12px;color:#8b949e;font-family:'DM Mono',monospace;">{{ wasabi_arrived }}</span>{% endif %}
            </div>
            <div style="font-size:12px;color:#484f58;margin-bottom:14px;">Choose how to deliver this disclosure record:</div>
            <div style="display:flex;flex-direction:column;gap:10px;">
                <label style="display:flex;align-items:center;gap:10px;cursor:pointer;margin:0;text-transform:none;letter-spacing:0;font-size:14px;color:#e6edf3;font-weight:400;">
                    <input type="radio" name="delivery" value="download" style="width:auto;"> &#11015;&#65039; Download only
                </label>
                <label style="display:flex;align-items:center;gap:10px;cursor:pointer;margin:0;text-transform:none;letter-spacing:0;font-size:14px;color:#e6edf3;font-weight:400;">
                    <input type="radio" name="delivery" value="cloud" style="width:auto;"> &#9729;&#65039; Save into cloud package only
                </label>
                <label style="display:flex;align-items:center;gap:10px;cursor:pointer;margin:0;text-transform:none;letter-spacing:0;font-size:14px;color:#e6edf3;font-weight:400;">
                    <input type="radio" name="delivery" value="both" checked style="width:auto;"> &#11015;&#65039;&#9729;&#65039; Download + save into cloud package
                </label>
            </div>
            <div style="font-size:11px;color:#484f58;margin-top:12px;line-height:1.5;">Saving into the cloud package retains this bookmark's footage beyond the normal overwrite cycle under the Wasabi retention workflow.</div>
            <input type="hidden" name="wasabi_prefix" value="{{ wasabi_prefix }}">
        </div>
        {% else %}
        <div style="background:#161b22;border:1px solid #30363d;border-radius:8px;padding:12px 16px;margin-bottom:14px;">
            {% if '[D]' in (bm.name or '').upper() %}
            <div style="font-size:12px;color:#8b949e;">&#8987; <strong style="color:#d29922;">[D] Deferred upload</strong> — footage not yet in cloud. Check again after 06:00. <strong style="color:#e6edf3;">Download only for now.</strong></div>
            {% elif '[C]' in (bm.name or '').upper() %}
            <div style="font-size:12px;color:#8b949e;">&#9888;&#65039; <strong style="color:#f85149;">[C] Cloud bookmark</strong> — footage not found in cloud. <strong style="color:#e6edf3;">Download only.</strong> Contact CCTV team if upload is expected.</div>
            {% else %}
            <div style="font-size:12px;color:#8b949e;">&#9729;&#65039; No footage found in cloud for this bookmark — <strong style="color:#e6edf3;">Download only</strong>.</div>
            {% endif %}
            <input type="hidden" name="delivery" value="download">
            <input type="hidden" name="wasabi_prefix" value="">
        </div>
        {% endif %}
        <!-- ───────────────────────────────────────────────────────────────── -->

        <button type="submit" class="submit-btn" id="foiSubmitBtn">&#128203; Generate Disclosure Record</button>

        <div id="loadingBox">
            <div style="font-size:36px;margin-bottom:12px;">⏳</div>
            <div style="font-size:16px;font-weight:700;color:#e6edf3;margin-bottom:8px;">Generating Disclosure Record...</div>
            <div style="font-size:13px;color:#8b949e;">Building your document — this is instant.</div>
        </div>

    </form>
</div>

<div id="successBox" style="display:none;max-width:820px;margin:0 auto;padding:0 20px 80px;">
    <div style="background:#120a1f;border:1px solid #8a5cf6;border-radius:10px;padding:30px;text-align:center;">
        <div style="font-size:44px;margin-bottom:12px;">&#10003;</div>
        <div style="font-size:20px;font-weight:700;color:#8a5cf6;margin-bottom:6px;">Disclosure Record Generated</div>
        <div id="foiSuccessSubtitle" style="font-size:13px;color:#8b949e;margin-bottom:16px;">Your document is downloading now.</div>
        <div id="foiCloudTimeline" style="display:none;background:#0d1117;border:1px solid #21262d;border-radius:8px;padding:14px 18px;margin-bottom:20px;text-align:left;">
            <div style="font-size:10px;color:#484f58;text-transform:uppercase;letter-spacing:1.2px;margin-bottom:10px;font-family:'DM Mono',monospace;">Cloud Package Timeline</div>
            <div style="display:flex;flex-direction:column;gap:9px;">
                <div style="display:flex;gap:12px;align-items:baseline;">
                    <span style="font-size:11px;color:#484f58;font-family:'DM Mono',monospace;min-width:160px;flex-shrink:0;">&#128230; Package confirmed</span>
                    <span id="foiTlArrived" style="font-size:13px;color:#8b949e;font-family:'DM Mono',monospace;"></span>
                </div>
                <div style="display:flex;gap:12px;align-items:baseline;">
                    <span style="font-size:11px;color:#484f58;font-family:'DM Mono',monospace;min-width:160px;flex-shrink:0;">&#128196; Record saved</span>
                    <span id="foiTlSaved" style="font-size:13px;color:#8a5cf6;font-family:'DM Mono',monospace;"></span>
                </div>
            </div>
        </div>
        <div id="foiCloudError" style="display:none;font-size:13px;color:#f85149;margin-bottom:20px;"></div>
        <a href="/" style="background:#21262d;color:#e6edf3;padding:14px 22px;border-radius:6px;font-size:14px;font-weight:600;text-decoration:none;display:inline-block;">&#8592; Back to Bookmarks</a>
    </div>
    <div style="background:#161b22;border:1px solid #30363d;border-radius:10px;padding:24px;margin-top:12px;">
        <div style="font-size:13px;font-weight:700;color:#8b949e;text-transform:uppercase;letter-spacing:1px;margin-bottom:16px;padding-bottom:10px;border-bottom:1px solid #21262d;">Email Disclosure Record</div>
        <div style="display:grid;grid-template-columns:1fr auto 1fr;gap:8px;align-items:end;">
            <div>
                <div style="font-size:11px;color:#8b949e;margin-bottom:6px;text-transform:uppercase;letter-spacing:0.4px;">Recipient Name</div>
                <input type="text" id="foiEmailRecipName" placeholder="e.g. john.smith" style="width:100%;padding:10px 14px;background:#0d1117;border:1px solid #30363d;border-radius:6px;font-size:14px;color:#e6edf3;font-family:'DM Sans',sans-serif;">
            </div>
            <div style="padding-bottom:2px;font-size:18px;color:#484f58;text-align:center;">@</div>
            <div>
                <div style="font-size:11px;color:#8b949e;margin-bottom:6px;text-transform:uppercase;letter-spacing:0.4px;">Domain</div>
                <select id="foiEmailDomain" style="width:100%;padding:10px 14px;background:#0d1117;border:1px solid #30363d;border-radius:6px;font-size:14px;color:#e6edf3;font-family:'DM Sans',sans-serif;">
                    <option value="rotherham.gov.uk">rotherham.gov.uk</option>
                    <option value="southyorkshire.police.uk">southyorkshire.police.uk</option>
                </select>
            </div>
        </div>
        <div id="foiEmailStatus" style="font-size:13px;margin-top:10px;min-height:20px;"></div>
        <button onclick="sendFoiEmail()" id="sendFoiEmailBtn" style="width:100%;margin-top:12px;padding:12px;background:#2a1a4a;color:#8a5cf6;border:1px solid #6e40c9;border-radius:6px;font-size:14px;font-weight:600;cursor:pointer;font-family:'DM Sans',sans-serif;">
            Send Disclosure Record by Email
        </button>
    </div>
</div>
<script>
var _wasabiArrived = "{{ wasabi_arrived }}";
// ── FOI clock ─────────────────────────────────────────────────────────────────
function foiGetUKTime() {
    var now = new Date();
    function lastSun(y,m){var d=new Date(Date.UTC(y,m,31));d.setUTCDate(31-d.getUTCDay());return d;}
    var y=now.getUTCFullYear();
    var isBST=now>=lastSun(y,2)&&now<lastSun(y,9);
    var uk=new Date(now.getTime()+(isBST?1:0)*3600000);
    var hh=String(uk.getUTCHours()).padStart(2,'0');
    var mm=String(uk.getUTCMinutes()).padStart(2,'0');
    var ss=String(uk.getUTCSeconds()).padStart(2,'0');
    var months=['January','February','March','April','May','June','July','August','September','October','November','December'];
    var ld=document.getElementById('foiLiveDate');
    var lt=document.getElementById('foiLiveTime');
    var tz=document.getElementById('foiTzLabel');
    if(ld) ld.textContent=uk.getUTCDate()+' '+months[uk.getUTCMonth()]+' '+uk.getUTCFullYear();
    if(lt) lt.textContent=hh+':'+mm+':'+ss;
    if(tz) tz.textContent=(isBST?'BST':'GMT')+' · Updates every second';
}
foiGetUKTime();
setInterval(foiGetUKTime, 1000);

function foiFetchServerTime(){
    fetch('/server-time').then(function(r){return r.json();}).then(function(d){
        var sd=document.getElementById('foiServerDate');
        var sh=document.getElementById('foiServerHHMM');
        var ss=document.getElementById('foiServerSS');
        var df=document.getElementById('foiServerDiff');
        var lt=document.getElementById('foiLiveTime');
        if(sd) sd.textContent=d.date;
        if(sh) sh.textContent=d.time.split(':')[0]+':'+d.time.split(':')[1]+':';
        if(ss) ss.textContent=d.time.split(':')[2];
        if(lt && df){
            var bt=lt.textContent;
            if(bt && bt!=='--:--:--'){
                var b=bt.split(':').map(Number);
                var s=d.time.split(':').map(Number);
                var diff=(s[0]*3600+s[1]*60+s[2])-(b[0]*3600+b[1]*60+b[2]);
                var abs=Math.abs(diff);
                var dm=Math.floor(abs/60),ds=abs%60;
                df.textContent=abs<10?'Within 10 seconds — accurate':(dm>0?'Server '+(diff>0?'fast':'slow')+' by '+dm+'m '+ds+'s':'Server '+(diff>0?'fast':'slow')+' by '+ds+'s');
            }
        }
    }).catch(function(){
        var df=document.getElementById('foiServerDiff');
        if(df) df.textContent='Could not reach server';
    });
}
foiFetchServerTime();
setInterval(foiFetchServerTime, 5000);

// ── FOI sync button ───────────────────────────────────────────────────────────
function foiUseTheseTimes() {
    var btTime  = document.getElementById('foiLiveTime').textContent;
    var hhmm    = document.getElementById('foiServerHHMM').textContent;
    var secPart = document.getElementById('foiServerSS').textContent;
    var srvTime = hhmm + secPart;
    if (btTime === '--:--:--' || srvTime === '--:--:--') {
        alert('Times not yet loaded — wait a moment and try again.');
        return;
    }
    var cb = document.getElementById('timeVerifyCheck');
    if (cb) { cb.checked = true; toggleTimeVerify(cb); }
    var bf = document.querySelector('[name="foi_verify_time"]');
    var nf = document.querySelector('[name="foi_system_time"]');
    var df = document.querySelector('[name="foi_time_offset"]');
    if (bf) bf.value = btTime;
    if (nf) nf.value = srvTime;
    var b = btTime.split(':').map(Number);
    var s = srvTime.split(':').map(Number);
    var diffSecs = (s[0]*3600+s[1]*60+s[2]) - (b[0]*3600+b[1]*60+b[2]);
    var absDiff  = Math.abs(diffSecs);
    var dMins = Math.floor(absDiff / 60);
    var dSecs = absDiff % 60;
    if (df) {
        var dir = absDiff === 0 ? '' : (' ' + (diffSecs > 0 ? 'fast' : 'slow'));
        if (absDiff === 0)     df.value = 'accurate';
        else if (dMins > 0)    df.value = dMins + ' minute' + (dMins !== 1 ? 's' : '') + ' and ' + dSecs + ' second' + (dSecs !== 1 ? 's' : '') + dir;
        else                   df.value = absDiff + ' second' + (absDiff !== 1 ? 's' : '') + dir;
    }
    var gs = document.querySelector('.green-section');
    if (gs) gs.scrollIntoView({behavior:'smooth', block:'center'});
}

// ── Legal basis badge ─────────────────────────────────────────────────────────
function updateLegalBasis() {
    var typeEl = document.getElementById('foiRequestType');
    var identYes = document.getElementById('identifiableYes');
    var badge = document.getElementById('legalBadge');
    var warning = document.getElementById('identifiableWarning');
    var display = document.getElementById('legalDisplay');
    if(!typeEl || !identYes) return;
    var type = typeEl.value;
    var identifiable = identYes.checked;
    if(warning) warning.style.display = identifiable ? 'block' : 'none';
    var text = '';
    if (identifiable) {
        text = 'UK GDPR Art.6(1)(f) / DPA 2018 — Data Protection';
    } else if (type === 'Public') {
        text = 'FOIA 2000';
    } else if (type === 'Solicitor' || type === 'Insurance') {
        text = 'DPA 2018 Schedule 2 — Legal Claims';
    } else {
        text = 'UK GDPR Art.6(1)(c) — Legal Obligation';
    }
    if(badge) badge.textContent = text;
    if(display) display.value = text;
}

// ── Time verify toggle ────────────────────────────────────────────────────────
function toggleTimeVerify(cb) {
    var f = document.getElementById('timeVerifyFields');
    if (f) f.style.display = cb.checked ? 'block' : 'none';
}

// ── Form submit ───────────────────────────────────────────────────────────────
var foiForm = document.getElementById('foiForm');
var _foiDocToken = null;
if (foiForm) foiForm.addEventListener('submit', function(e) {
    e.preventDefault();
    var btn = document.getElementById('foiSubmitBtn');
    btn.disabled = true; btn.textContent = '⏳ Generating...';
    document.getElementById('loadingBox').style.display = 'block';
    fetch(window.location.href, {method:'POST', body: new FormData(this)})
        .then(function(r) {
            if (!r.ok) throw new Error('Server error ' + r.status);
            _foiDocToken = r.headers.get('X-Doc-Token');
            window._foiDelivery      = r.headers.get('X-Delivery') || 'download';
            window._foiWasabiUploaded = r.headers.get('X-Wasabi-Uploaded') === 'true';
            window._foiWasabiTs      = r.headers.get('X-Wasabi-Timestamp') || '';
            window._foiWasabiErr     = r.headers.get('X-Wasabi-Error') || '';
            return r.blob();
        })
        .then(function(blob) {
            var delivery = window._foiDelivery || 'download';
            if (delivery === 'download' || delivery === 'both') {
                var url = URL.createObjectURL(blob);
                var a = document.createElement('a'); a.href = url;
                a.download = 'foi_disclosure_record.docx';
                document.body.appendChild(a); a.click();
                document.body.removeChild(a); URL.revokeObjectURL(url);
            }
            document.getElementById('loadingBox').style.display = 'none';
            document.querySelector('.container').style.display  = 'none';
            document.getElementById('successBox').style.display = 'block';
            var sub = document.getElementById('foiSuccessSubtitle');
            if (sub) {
                if (delivery === 'cloud') sub.textContent = 'Record saved into cloud package — no local download.';
                else if (delivery === 'both') sub.textContent = 'Record downloading and saved into cloud package.';
                else sub.textContent = 'Your document is downloading now.';
            }
            // Cloud package timeline
            if (window._foiWasabiUploaded) {
                var tl = document.getElementById('foiCloudTimeline');
                if (tl) {
                    document.getElementById('foiTlArrived').textContent = _wasabiArrived || '—';
                    document.getElementById('foiTlSaved').textContent   = window._foiWasabiTs || '—';
                    tl.style.display = 'block';
                }
            } else if (window._foiWasabiErr) {
                var ce = document.getElementById('foiCloudError');
                if (ce) {
                    ce.textContent = '⚠ Record upload to cloud failed: ' + window._foiWasabiErr;
                    ce.style.display = 'block';
                }
            }
        })
        .catch(function(err) {
            document.getElementById('loadingBox').innerHTML =
                '<div style="color:#f85149;font-size:15px;margin-bottom:12px;">❌ ' + err.message + '</div>' +
                '<button onclick="location.reload()" style="background:#21262d;color:#e6edf3;padding:10px 20px;border-radius:6px;border:none;cursor:pointer;">Try Again</button>';
        });
});

function sendFoiEmail() {
    var name   = document.getElementById('foiEmailRecipName').value.trim();
    var domain = document.getElementById('foiEmailDomain').value;
    var status = document.getElementById('foiEmailStatus');
    var btn    = document.getElementById('sendFoiEmailBtn');
    if (!name) { status.style.color='#f85149'; status.textContent='Please enter a recipient name.'; return; }
    if (!_foiDocToken) { status.style.color='#f85149'; status.textContent='No document token — please regenerate the record.'; return; }
    btn.disabled = true; btn.textContent = '⏳ Sending...';
    status.style.color = '#8b949e'; status.textContent = 'Sending to ' + name + '@' + domain + '...';
    fetch('/send-email', {
        method: 'POST',
        headers: {'Content-Type': 'application/json'},
        body: JSON.stringify({token: _foiDocToken, recipient_name: name, recipient_domain: domain})
    })
    .then(function(r) { return r.json(); })
    .then(function(data) {
        if (data.ok) {
            status.style.color = '#8a5cf6';
            status.textContent = '✅ Email sent successfully to ' + data.recipient;
            btn.textContent = 'Send Another Copy';
            btn.disabled = false;
        } else {
            status.style.color = '#f85149';
            status.textContent = '❌ ' + data.error;
            btn.textContent = 'Send Disclosure Record by Email';
            btn.disabled = false;
        }
    })
    .catch(function(err) {
        status.style.color = '#f85149';
        status.textContent = '❌ Network error: ' + err.message;
        btn.textContent = 'Send Disclosure Record by Email';
        btn.disabled = false;
    });
}
</script>
</body></html>"""
# ── Routes ────────────────────────────────────────────────────────────────────

@app.route("/login", methods=["GET","POST"])
def login():
    error = None
    if request.method == "POST":
        username = request.form.get("username","").strip().lower()
        password = request.form.get("password","")
        if check_password(username, password):
            session["username"]  = username
            session["user_name"] = USERS[username]["name"]
            session["user_role"] = USERS[username]["role"]
            session["initials"]  = USERS[username]["initials"]
            return redirect(url_for("index"))
        error = "Invalid username or password."
    return render_template_string(LOGIN_HTML, error=error, site_ref=SITE_REF)

@app.route("/logout")
def logout():
    session.clear(); return redirect(url_for("login"))

@app.route("/")
@login_required
def index():
    bookmarks = get_bookmarks()
    # Run Wasabi check for every bookmark so the index page shows real cloud status
    for bm in bookmarks:
        try:
            found, _, arrived = check_wasabi_footage(bm.get("name", ""))
            bm["wasabi_confirmed"] = found
            bm["wasabi_arrived"]   = arrived or ""
        except Exception:
            bm["wasabi_confirmed"] = False
            bm["wasabi_arrived"]   = ""
    return render_template_string(BOOKMARKS_HTML, bookmarks=bookmarks, session=session, site_ref=SITE_REF)

@app.route("/form/<int:bookmark_id>")
@login_required
def form(bookmark_id):
    return redirect(url_for("syp", bookmark_id=bookmark_id))

@app.route("/syp/<int:bookmark_id>", methods=["GET","POST"])
@login_required
def syp(bookmark_id):
    return _statement_handler(bookmark_id, pipeline="syp")

@app.route("/rmbc/<int:bookmark_id>", methods=["GET","POST"])
@login_required
def rmbc(bookmark_id):
    return _statement_handler(bookmark_id, pipeline="rmbc")

def _statement_handler(bookmark_id, pipeline):
    bm = get_bookmark(bookmark_id)
    if not bm: return redirect(url_for("index"))

    # ── Check Wasabi for matching footage on every GET ────────────────────────
    wasabi_found   = False
    wasabi_prefix  = ""
    wasabi_arrived = ""   # HH:MM on DD/MM/YYYY — when cloud package became available
    wasabi_error   = ""
    try:
        wasabi_found, wasabi_prefix, wasabi_arrived = check_wasabi_footage(bm.get("name", ""))
        wasabi_prefix  = wasabi_prefix  or ""
        wasabi_arrived = wasabi_arrived or ""
        if not wasabi_found:
            wasabi_error = "deferred"   # footage not yet uploaded
    except Exception as we:
        wasabi_error = str(we)
    # ─────────────────────────────────────────────────────────────────────────

    if request.method == "POST":
        form_data     = request.form.to_dict()
        form_data["_pipeline"] = pipeline          # passed to generate_statement for pipeline-aware wording
        delivery      = form_data.get("delivery", "download")  # download | cloud | both
        wp            = form_data.get("wasabi_prefix", "")
        download_time = datetime.now(timezone.utc)
        try:
            statement_text = generate_statement(bm, form_data, download_time)
            docx_buf       = build_docx(statement_text, form_data, download_time)
            witness_name   = form_data.get("witness_name", "Officer")
            bookmark_name  = bm.get("name", "statement")
            date_str  = download_time.strftime("%Y-%m-%d_%H-%M")
            safe_name = "".join(c if c.isalnum() or c in "-_" else "_" for c in witness_name)
            safe_bm   = "".join(c if c.isalnum() or c in "-_" else "_" for c in bookmark_name)
            pipeline_tag  = "SYP" if pipeline == "syp" else "RMBC"
            filename       = f"{date_str}_{pipeline_tag}_{safe_name}_{safe_bm}.docx"
            cloud_filename = f"Statement-{SITE_REF}-{download_time.strftime('%d-%m-%Y-%H-%M')}.docx"
            # Save to temp for email feature
            doc_token = str(uuid.uuid4())
            tmp_path  = os.path.join(tempfile.gettempdir(), f"{doc_token}.docx")
            with open(tmp_path, "wb") as f:
                f.write(docx_buf.getvalue())
            _TEMP_DOCS[doc_token] = (tmp_path, filename)
            # ── Wasabi upload if delivery includes cloud ──────────────────────
            wasabi_ts  = ""
            wasabi_err = ""
            if delivery in ("cloud", "both") and wp:
                try:
                    docx_buf.seek(0)
                    wasabi_ts, _ = upload_to_wasabi(docx_buf.read(), wp, cloud_filename)
                except Exception as ue:
                    wasabi_err = str(ue)
                    print(f"Wasabi upload error: {ue}")
            # ─────────────────────────────────────────────────────────────────
            docx_buf.seek(0)
            resp = send_file(docx_buf, as_attachment=True, download_name=filename,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            resp.headers["X-Doc-Token"]    = doc_token
            resp.headers["X-Delivery"]     = delivery
            resp.headers["X-Wasabi-Uploaded"]  = "true" if wasabi_ts else "false"
            resp.headers["X-Wasabi-Timestamp"] = wasabi_ts
            resp.headers["X-Wasabi-Error"]     = wasabi_err
            resp.headers["Access-Control-Expose-Headers"] = (
                "X-Doc-Token, X-Delivery, X-Wasabi-Uploaded, X-Wasabi-Timestamp, X-Wasabi-Error"
            )
            return resp
        except Exception as e:
            import traceback; traceback.print_exc()
            route = "syp" if pipeline == "syp" else "rmbc"
            return (f"<html><body style='font-family:Arial;padding:40px;background:#0d1117;color:#f85149;'>"
                    f"<h2>Error: {str(e)}</h2>"
                    f"<pre style='color:#f85149;font-size:12px;margin-top:20px;white-space:pre-wrap;'>{traceback.format_exc()}</pre>"
                    f"<a href='/{route}/{bookmark_id}' style='color:#58a6ff;'>← Try again</a>"
                    f"</body></html>"), 500
    today    = datetime.now().strftime("%d/%m/%Y")
    initials = session.get("initials", "XX")
    return render_template_string(FORM_HTML, bm=bm, session=session, today=today,
                                  locations=LOCATIONS, initials=initials, pipeline=pipeline,
                                  wasabi_found=wasabi_found, wasabi_prefix=wasabi_prefix,
                                  wasabi_arrived=wasabi_arrived)

@app.route("/foi/<int:bookmark_id>", methods=["GET","POST"])
@login_required
def foi(bookmark_id):
    bm = get_bookmark(bookmark_id)
    if not bm: return redirect(url_for("index"))

    # ── Check Wasabi on GET ───────────────────────────────────────────────────
    wasabi_found   = False
    wasabi_prefix  = ""
    wasabi_arrived = ""
    try:
        wasabi_found, wasabi_prefix, wasabi_arrived = check_wasabi_footage(bm.get("name", ""))
        wasabi_prefix  = wasabi_prefix  or ""
        wasabi_arrived = wasabi_arrived or ""
    except Exception as we:
        print(f"Wasabi check error (FOI): {we}")
    # ─────────────────────────────────────────────────────────────────────────

    if request.method == "POST":
        form_data     = request.form.to_dict()
        delivery      = form_data.get("delivery", "download")
        wp            = form_data.get("wasabi_prefix", "")
        download_time = datetime.now(timezone.utc)
        try:
            foi_data = assemble_foi_data(bm, form_data, download_time)
            docx_buf = build_foi_docx(foi_data, download_time)
            witness_name = form_data.get("witness_name", "Officer")
            date_str  = download_time.strftime("%Y-%m-%d_%H-%M")
            safe_name = "".join(c if c.isalnum() or c in "-_" else "_" for c in witness_name)
            filename       = f"{date_str}_{safe_name}_FOI_Disclosure.docx"
            cloud_filename = f"Statement-{SITE_REF}-{download_time.strftime('%d-%m-%Y-%H-%M')}.docx"
            doc_token = str(uuid.uuid4())
            tmp_path  = os.path.join(tempfile.gettempdir(), f"{doc_token}.docx")
            with open(tmp_path, "wb") as f:
                f.write(docx_buf.getvalue())
            _TEMP_DOCS[doc_token] = (tmp_path, filename)
            # ── Wasabi upload ─────────────────────────────────────────────────
            wasabi_ts  = ""
            wasabi_err = ""
            if delivery in ("cloud", "both") and wp:
                try:
                    docx_buf.seek(0)
                    wasabi_ts, _ = upload_to_wasabi(docx_buf.read(), wp, cloud_filename)
                except Exception as ue:
                    wasabi_err = str(ue)
                    print(f"Wasabi upload error (FOI): {ue}")
            # ─────────────────────────────────────────────────────────────────
            docx_buf.seek(0)
            resp = send_file(docx_buf, as_attachment=True, download_name=filename,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            resp.headers["X-Doc-Token"]        = doc_token
            resp.headers["X-Delivery"]         = delivery
            resp.headers["X-Wasabi-Uploaded"]  = "true" if wasabi_ts else "false"
            resp.headers["X-Wasabi-Timestamp"] = wasabi_ts
            resp.headers["X-Wasabi-Error"]     = wasabi_err
            resp.headers["Access-Control-Expose-Headers"] = (
                "X-Doc-Token, X-Delivery, X-Wasabi-Uploaded, X-Wasabi-Timestamp, X-Wasabi-Error"
            )
            return resp
        except Exception as e:
            import traceback; traceback.print_exc()
            return (f"<html><body style='font-family:Arial;padding:40px;background:#0d1117;color:#f85149;'>"
                    f"<h2>Error: {str(e)}</h2>"
                    f"<pre style='color:#f85149;font-size:12px;margin-top:20px;white-space:pre-wrap;'>{traceback.format_exc()}</pre>"
                    f"<a href='/foi/{bookmark_id}' style='color:#58a6ff;'>← Try again</a>"
                    f"</body></html>"), 500
    today    = datetime.now().strftime("%d/%m/%Y")
    initials = session.get("initials", "XX")
    return render_template_string(FOI_FORM_HTML, bm=bm, session=session,
                                  today=today, initials=initials, site_ref=SITE_REF,
                                  wasabi_found=wasabi_found, wasabi_prefix=wasabi_prefix,
                                  wasabi_arrived=wasabi_arrived)

@app.route("/send-email", methods=["POST"])
@login_required
def send_email():
    try:
        data       = request.get_json()
        token      = data.get("token", "")
        recip_name = data.get("recipient_name", "").strip()
        recip_dom  = data.get("recipient_domain", "rotherham.gov.uk")
        recipient  = f"{recip_name}@{recip_dom}"

        if not token or token not in _TEMP_DOCS:
            return jsonify({"ok": False, "error": "Document not found — please regenerate the statement first."})
        if not recip_name:
            return jsonify({"ok": False, "error": "Please enter a recipient name."})
        if not GMAIL_USER or not GMAIL_APP_PASSWORD:
            return jsonify({"ok": False, "error": "Email not configured on this server. Add GMAIL_USER and GMAIL_APP_PASSWORD to .env"})

        tmp_path, filename = _TEMP_DOCS[token]
        if not os.path.exists(tmp_path):
            return jsonify({"ok": False, "error": "Temporary file expired — please regenerate the statement."})

        # Build email
        msg = MIMEMultipart()
        msg["From"]    = GMAIL_USER
        msg["To"]      = recipient
        msg["Subject"] = f"RMBC CCTV — Witness Statement — {SITE_REF}"

        body = (
            f"Please find attached the CCTV witness statement generated by RMBC CCTV Evidence Management.\n\n"
            f"This document has been prepared by {session.get('user_name', 'RMBC CCTV')} "
            f"and is marked OFFICIAL SENSITIVE.\n\n"
            f"System Reference: {SITE_REF}\n"
            f"Generated: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n\n"
            f"Please do not reply to this email. For queries contact the RMBC CCTV team directly."
        )
        msg.attach(MIMEText(body, "plain"))

        with open(tmp_path, "rb") as f:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(f.read())
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", f'attachment; filename="{filename}"')
        msg.attach(part)

        # Send via Gmail SMTP TLS
        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.ehlo()
            server.starttls()
            server.login(GMAIL_USER, GMAIL_APP_PASSWORD)
            server.sendmail(GMAIL_USER, recipient, msg.as_string())

        # Clean up temp file after send
        try:
            os.remove(tmp_path)
            del _TEMP_DOCS[token]
        except Exception:
            pass

        return jsonify({"ok": True, "recipient": recipient})

    except smtplib.SMTPAuthenticationError:
        return jsonify({"ok": False, "error": "Gmail authentication failed. Check GMAIL_USER and GMAIL_APP_PASSWORD in .env"})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)})

@app.route("/server-time")
def server_time():
    now = datetime.now()
    return jsonify({"date": now.strftime("%d/%m/%Y"), "time": now.strftime("%H:%M:%S")})

if __name__ == "__main__":
    print(f"\n🤖  AI backend: Ollama ({OLLAMA_HOST}) — model: {OLLAMA_MODEL}")
    print(f"\n🎥 RMBC CCTV Evidence Management — {SITE_REF}")
    print("   http://0.0.0.0:5000\n")
    app.run(host="0.0.0.0", port=5000, debug=False)
